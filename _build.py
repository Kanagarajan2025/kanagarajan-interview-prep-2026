from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

H1_COLOR = RGBColor(0x1F, 0x49, 0x7D)
H2_COLOR = RGBColor(0x2E, 0x74, 0xB5)
H3_COLOR = RGBColor(0x1F, 0x6B, 0x96)
TBL_HEADER = 'DEEAF1'
CODE_BG    = 'F2F2F2'
BQ_BORDER  = '2E74B5'

def set_cell_shading(cell, fill):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),fill)
    tcPr.append(shd)

def set_para_shading(para, fill):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),fill)
    pPr.append(shd)

def set_para_border(para):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'),'single'); left.set(qn('w:sz'),'12')
    left.set(qn('w:space'),'8');    left.set(qn('w:color'),BQ_BORDER)
    pBdr.append(left); pPr.append(pBdr)

def set_hr(para):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'),'single'); bot.set(qn('w:sz'),'6')
    bot.set(qn('w:space'),'1');    bot.set(qn('w:color'),'BBBBBB')
    pBdr.append(bot); pPr.append(pBdr)

def add_hyperlink(para, url, text):
    part = para.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    hyperlink.set(qn('w:history'), '1')
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    color = OxmlElement('w:color'); color.set(qn('w:val'), '0563C1')
    u = OxmlElement('w:u'); u.set(qn('w:val'), 'single')
    rPr.append(color); rPr.append(u)
    r.append(rPr)
    t = OxmlElement('w:t'); t.text = text
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    r.append(t)
    hyperlink.append(r)
    para._p.append(hyperlink)

URL_PATTERN = re.compile(r'(https?://[^\s\)\]>,]+)')
INLINE = re.compile(r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)')

def add_runs(para, text, mono=False):
    if mono:
        run = para.add_run(text)
        run.font.name = 'Courier New'; run.font.size = Pt(8.5)
        return
    for part in URL_PATTERN.split(text):
        if not part: continue
        if URL_PATTERN.match(part):
            add_hyperlink(para, part, part)
        else:
            for chunk in INLINE.split(part):
                if not chunk: continue
                if chunk.startswith('**') and chunk.endswith('**'):
                    r = para.add_run(chunk[2:-2]); r.bold = True
                elif chunk.startswith('*') and chunk.endswith('*'):
                    r = para.add_run(chunk[1:-1]); r.italic = True
                elif chunk.startswith('`') and chunk.endswith('`'):
                    r = para.add_run(chunk[1:-1])
                    r.font.name = 'Courier New'; r.font.size = Pt(9)
                else:
                    para.add_run(chunk)

def split_row(line):
    safe = line.replace('\\|', '\x00')
    return [c.replace('\x00','|').strip() for c in safe.split('|')[1:-1]]

def build(md_path, out_path):
    with open(md_path, encoding='utf-8') as f:
        lines = [l.rstrip('\n') for l in f]

    doc = Document()
    for sec in doc.sections:
        sec.top_margin=Inches(0.9); sec.bottom_margin=Inches(0.9)
        sec.left_margin=Inches(1.1); sec.right_margin=Inches(1.1)
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(10.5)

    i = 0; in_code = False; code_buf = []
    in_table = False; tbl_rows = []

    def flush_table():
        nonlocal in_table, tbl_rows
        if not tbl_rows: in_table=False; return
        cols = max(len(r) for r in tbl_rows)
        tbl = doc.add_table(rows=len(tbl_rows), cols=cols)
        tbl.style = 'Table Grid'
        for ri, rd in enumerate(tbl_rows):
            row = tbl.rows[ri]
            for ci in range(cols):
                ct = rd[ci] if ci < len(rd) else ''
                cell = row.cells[ci]
                p = cell.paragraphs[0]; p.clear()
                p.paragraph_format.space_before=Pt(1); p.paragraph_format.space_after=Pt(1)
                add_runs(p, ct)
                if ri == 0:
                    for run in p.runs: run.bold=True
                    set_cell_shading(cell, TBL_HEADER)
        tbl.autofit = True
        doc.add_paragraph()
        in_table=False; tbl_rows.clear()

    def flush_code():
        nonlocal code_buf
        if not code_buf: code_buf=[]; return
        p = doc.add_paragraph()
        p.paragraph_format.left_indent=Inches(0.25)
        p.paragraph_format.space_before=Pt(2); p.paragraph_format.space_after=Pt(4)
        set_para_shading(p, CODE_BG)
        add_runs(p, '\n'.join(code_buf), mono=True)
        code_buf=[]

    while i < len(lines):
        line = lines[i]

        if line.startswith('```'):
            if in_table: flush_table()
            if in_code: flush_code(); in_code=False
            else: in_code=True
            i+=1; continue

        if in_code:
            code_buf.append(line); i+=1; continue

        if line.startswith('|'):
            if re.match(r'\|[\s\-:|]+\|', line): i+=1; continue
            cells = split_row(line)
            nxt = lines[i+1] if i+1<len(lines) else ''
            if re.match(r'\|[\s\-:|]+\|', nxt):
                if not in_table: in_table=True; tbl_rows.clear()
                tbl_rows.append(cells); i+=2; continue
            else:
                in_table=True; tbl_rows.append(cells); i+=1; continue
        else:
            if in_table: flush_table()

        if line.strip() in ('---','***','___'):
            p=doc.add_paragraph()
            p.paragraph_format.space_before=Pt(2); p.paragraph_format.space_after=Pt(2)
            set_hr(p); i+=1; continue

        m = re.match(r'^(#{1,4}) (.+)', line)
        if m:
            lvl=len(m.group(1))
            txt=re.sub(r'\*\*([^*]+)\*\*',r'\1',m.group(2))
            txt=re.sub(r'\*([^*]+)\*',r'\1',txt)
            h=doc.add_heading(txt, level=lvl)
            color=[H1_COLOR,H2_COLOR,H3_COLOR,H3_COLOR][lvl-1]
            for run in h.runs: run.font.color.rgb=color
            i+=1; continue

        if line.startswith('> '):
            text=line[2:].strip()
            p=doc.add_paragraph()
            p.paragraph_format.left_indent=Inches(0.35)
            p.paragraph_format.space_before=Pt(1); p.paragraph_format.space_after=Pt(1)
            set_para_border(p); add_runs(p, text)
            i+=1; continue

        if re.match(r'^[-*] ', line):
            p=doc.add_paragraph(style='List Bullet'); add_runs(p, line[2:]); i+=1; continue

        if re.match(r'^\d+\. ', line):
            p=doc.add_paragraph(style='List Number')
            add_runs(p, re.sub(r'^\d+\. ','',line)); i+=1; continue

        if line.strip()=='':
            doc.add_paragraph(); i+=1; continue

        p=doc.add_paragraph(); add_runs(p, line.rstrip()); i+=1

    if in_table: flush_table()
    if in_code:  flush_code()

    doc.save(out_path)
    print(f'Saved: {out_path}')

build('Day_01_Linux_Basics.md', 'Day_01_Linux_Basics_FINAL_new.docx')

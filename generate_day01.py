from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Page margins ───────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin   = Cm(2.0)
    section.right_margin  = Cm(2.0)

# ── Colour palette ─────────────────────────────────────────────────────────────
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
BLACK      = RGBColor(0x00, 0x00, 0x00)
NAVY       = RGBColor(0x1A, 0x23, 0x3E)
BLUE       = RGBColor(0x1F, 0x6F, 0xEB)
ORANGE     = RGBColor(0xE0, 0x6C, 0x00)
GREEN      = RGBColor(0x1B, 0x7A, 0x3E)
RED        = RGBColor(0xCC, 0x00, 0x00)
GREY_TEXT  = RGBColor(0x88, 0x88, 0x99)
CODE_FG    = RGBColor(0xCD, 0xD6, 0xF4)
COMMENT_FG = RGBColor(0x6C, 0x91, 0xC2)

# ══════════════════════════════════════════════════════════════════════════════
#  HELPER FUNCTIONS
# ══════════════════════════════════════════════════════════════════════════════

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def set_cell_borders(cell, color='CCCCCC'):
    tc     = cell._tc
    tcPr   = tc.get_or_add_tcPr()
    tcBdr  = OxmlElement('w:tcBorders')
    for side in ('top','left','bottom','right'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'),   'single')
        el.set(qn('w:sz'),    '4')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), color)
        tcBdr.append(el)
    tcPr.append(tcBdr)

def heading1(text):
    tbl  = doc.add_table(rows=1, cols=1)
    tbl.style = 'Table Grid'
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, '1A233E')
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after  = Pt(5)
    p.paragraph_format.left_indent  = Pt(8)
    r = p.add_run(text.upper())
    r.bold = True; r.font.color.rgb = WHITE
    r.font.size = Pt(13); r.font.name = 'Calibri'
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(2)

def heading2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(3)
    r = p.add_run('● ' + text)
    r.bold = True; r.font.color.rgb = BLUE
    r.font.size = Pt(12); r.font.name = 'Calibri'

def heading3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run('▸ ' + text)
    r.bold = True; r.font.color.rgb = ORANGE
    r.font.size = Pt(11); r.font.name = 'Calibri'

def body(text, indent=0, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.left_indent  = Pt(indent)
    r = p.add_run(text)
    r.font.size = Pt(10.5); r.font.name = 'Calibri'
    r.font.color.rgb = BLACK

def bullet(text, indent=18, char='•'):
    p = doc.add_paragraph()
    p.paragraph_format.space_before      = Pt(0)
    p.paragraph_format.space_after       = Pt(3)
    p.paragraph_format.left_indent       = Pt(indent)
    p.paragraph_format.first_line_indent = Pt(-10)
    r = p.add_run(f'{char}  {text}')
    r.font.size = Pt(10.5); r.font.name = 'Calibri'; r.font.color.rgb = BLACK

def sub_bullet(text):
    bullet(text, indent=34, char='◦')

def important_box(text, label='IMPORTANT'):
    tbl  = doc.add_table(rows=1, cols=1)
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, 'FFF3CD')
    set_cell_borders(cell, 'E0A800')
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after  = Pt(5)
    p.paragraph_format.left_indent  = Pt(6)
    r1 = p.add_run(f'⚠  {label}:  ')
    r1.bold = True; r1.font.color.rgb = ORANGE
    r1.font.size = Pt(10.5); r1.font.name = 'Calibri'
    r2 = p.add_run(text)
    r2.font.color.rgb = RGBColor(0x4A,0x3A,0x00)
    r2.font.size = Pt(10.5); r2.font.name = 'Calibri'
    sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(3)

def tip_box(text, label='TIP'):
    tbl  = doc.add_table(rows=1, cols=1)
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, 'D4EDDA')
    set_cell_borders(cell, '28A745')
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after  = Pt(5)
    p.paragraph_format.left_indent  = Pt(6)
    r1 = p.add_run(f'✓  {label}:  ')
    r1.bold = True; r1.font.color.rgb = GREEN
    r1.font.size = Pt(10.5); r1.font.name = 'Calibri'
    r2 = p.add_run(text)
    r2.font.color.rgb = RGBColor(0x15,0x4C,0x2A)
    r2.font.size = Pt(10.5); r2.font.name = 'Calibri'
    sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(3)

def youtube_box(videos):
    """videos = list of (topic, channel, url) tuples"""
    tbl = doc.add_table(rows=1 + len(videos), cols=3)
    tbl.style = 'Table Grid'
    hdrs = ['Topic', 'Channel', 'YouTube Link']
    h_colors = ['CC0000', 'AA0000', 'BB0000']
    for i, (hdr, bg) in enumerate(zip(hdrs, h_colors)):
        cell = tbl.cell(0, i)
        set_cell_bg(cell, bg)
        set_cell_borders(cell, 'FFFFFF')
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        r = p.add_run(('▶ ' if i==0 else '') + hdr)
        r.bold = True; r.font.color.rgb = WHITE
        r.font.size = Pt(10); r.font.name = 'Calibri'
    for ri, (topic, channel, url) in enumerate(videos):
        shade = 'FFF5F5' if ri % 2 == 0 else 'FFECEC'
        for ci, val in enumerate([topic, channel, url]):
            cell = tbl.cell(ri+1, ci)
            set_cell_bg(cell, shade)
            set_cell_borders(cell, 'FFCCCC')
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            r = p.add_run(val)
            if ci == 2:
                r.font.color.rgb = BLUE
                r.underline = True
            else:
                r.font.color.rgb = BLACK
            r.font.size = Pt(9.5); r.font.name = 'Calibri'
    sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(4)

def code_block(lines):
    tbl  = doc.add_table(rows=1, cols=1)
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, '1E1E2E')
    set_cell_borders(cell, '444455')
    cell.paragraphs[0].paragraph_format.space_before = Pt(6)
    cell.paragraphs[0].paragraph_format.left_indent  = Pt(10)
    first = True
    for line in lines:
        if first:
            para = cell.paragraphs[0]; first = False
        else:
            para = cell.add_paragraph()
            para.paragraph_format.left_indent  = Pt(10)
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after  = Pt(0)
        is_comment = line.strip().startswith('#')
        r = para.add_run(line)
        r.font.color.rgb = COMMENT_FG if is_comment else CODE_FG
        r.font.name = 'Courier New'; r.font.size = Pt(9)
    cell.paragraphs[-1].paragraph_format.space_after = Pt(6)
    sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(2)

def simple_table(headers, rows, header_bg='1F6FEB'):
    tbl = doc.add_table(rows=1+len(rows), cols=len(headers))
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, h in enumerate(headers):
        cell = tbl.cell(0, i)
        set_cell_bg(cell, header_bg)
        set_cell_borders(cell, 'FFFFFF')
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        r = p.add_run(h)
        r.bold = True; r.font.color.rgb = WHITE
        r.font.size = Pt(10); r.font.name = 'Calibri'
    for ri, row in enumerate(rows):
        shade = 'F8F9FA' if ri % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate(row):
            cell = tbl.cell(ri+1, ci)
            set_cell_bg(cell, shade)
            set_cell_borders(cell, 'DDDDDD')
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            r = p.add_run(val)
            r.font.size = Pt(10); r.font.name = 'Calibri'; r.font.color.rgb = BLACK
    sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(4)

def qa_block(question, answer):
    tbl = doc.add_table(rows=2, cols=1)
    q_cell = tbl.cell(0, 0)
    set_cell_bg(q_cell, '1A233E')
    set_cell_borders(q_cell, '1A233E')
    p = q_cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Pt(6)
    r = p.add_run('Q:  ' + question)
    r.bold = True; r.font.color.rgb = WHITE
    r.font.size = Pt(10.5); r.font.name = 'Calibri'
    a_cell = tbl.cell(1, 0)
    set_cell_bg(a_cell, 'EEF4FF')
    set_cell_borders(a_cell, 'AABBDD')
    p2 = a_cell.paragraphs[0]
    p2.paragraph_format.space_before = Pt(5)
    p2.paragraph_format.space_after  = Pt(5)
    p2.paragraph_format.left_indent  = Pt(6)
    r2 = p2.add_run(answer)
    r2.font.color.rgb = NAVY; r2.font.size = Pt(10.5); r2.font.name = 'Calibri'
    sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(6)

def divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    r = p.add_run('─' * 110)
    r.font.size = Pt(7); r.font.color.rgb = RGBColor(0xCC,0xCC,0xCC)

def page_break():
    doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  COVER BLOCK
# ══════════════════════════════════════════════════════════════════════════════
tbl  = doc.add_table(rows=1, cols=1)
cell = tbl.cell(0, 0)
set_cell_bg(cell, '1A233E'); set_cell_borders(cell, '1A233E')

p = cell.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(18); p.paragraph_format.space_after = Pt(4)
r = p.add_run('30-DAY INTERVIEW PREPARATION  —  KANAGARAJAN M')
r.bold = True; r.font.color.rgb = WHITE; r.font.size = Pt(15); r.font.name = 'Calibri'

p2 = cell.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.paragraph_format.space_before = Pt(4); p2.paragraph_format.space_after = Pt(2)
r2 = p2.add_run('DAY 1  —  LINUX COMMAND LINE BASICS')
r2.bold = True; r2.font.color.rgb = RGBColor(0x90,0xC8,0xFF)
r2.font.size = Pt(14); r2.font.name = 'Calibri'

p3 = cell.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
p3.paragraph_format.space_before = Pt(2); p3.paragraph_format.space_after = Pt(6)
r3 = p3.add_run('May 17, 2026  |  7:00 AM – 4:30 PM  |  8 Hours  |  Target: 24 LPA')
r3.font.color.rgb = RGBColor(0xBB,0xCC,0xEE); r3.font.size = Pt(10); r3.font.name = 'Calibri'

p4 = cell.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
p4.paragraph_format.space_before = Pt(2); p4.paragraph_format.space_after = Pt(14)
r4 = p4.add_run('Linux Foundation  |  SSH  |  Permissions  |  Processes  |  Piping  |  vim  |  DSA: Two Sum')
r4.font.color.rgb = RGBColor(0x77,0xAA,0xDD); r4.font.size = Pt(9.5); r4.font.name = 'Calibri'

sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(6)

tip_box(
    'This is your foundation day. AWS, Docker, Kubernetes, CI/CD — all run on Linux. '
    'Every server skill you build in the next 29 days assumes you know this. '
    'Every command you practice tonight becomes muscle memory.',
    label='WHY DAY 1 MATTERS'
)

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 1 — TIMELINE
# ══════════════════════════════════════════════════════════════════════════════
heading1('Section 1 — Your Complete Day Timeline (7:00 AM – 4:30 PM)')
body('8 hours of focused study. Breaks are NOT counted. Follow this exactly.', space_after=6)

simple_table(
    ['Time', 'Block', 'Activity', 'Study Hrs'],
    [
        ['7:00 – 7:15 AM',   'Read',      'Scan this entire document once — know what is coming',                           '0.25 hr'],
        ['7:15 – 9:15 AM',   'Theory 1',  'Linux intro + file system tree + all essential commands (Sections 3–5)',         '2 hrs'],
        ['9:15 – 9:30 AM',   'Break',     'Water, walk, step away from screen — mandatory rest',                           '—'],
        ['9:30 – 11:00 AM',  'Theory 2',  'Permissions + process management + text editors + piping (Sections 6–9)',        '1.5 hrs'],
        ['11:00 – 12:00 PM', 'Lunch',     'Full break — leave your desk. Brain consolidates during rest.',                  '—'],
        ['12:00 – 1:00 PM',  'DSA',       'Two Sum (LeetCode #1) — brute force, then HashMap, trace through examples',      '1 hr'],
        ['1:00 – 2:00 PM',   'Q&A',       'Answer all 13 interview questions OUT LOUD — no reading, speak from memory',      '1 hr'],
        ['2:00 – 2:10 PM',   'Break',     'Short break before practicals',                                                  '—'],
        ['2:10 – 4:10 PM',   'Practical', 'Hands-on: Git Bash + all commands + file ops + permissions + piping labs',       '2 hrs'],
        ['4:10 – 4:30 PM',   'Wrap-up',   'Write notes in your own words + scan Day 2 document',                           '0.33 hr'],
    ],
    header_bg='1A233E'
)

important_box(
    'Total study = 8 hours. Day ENDS at 4:30 PM. Rest after that. '
    'Do not run study blocks together — the breaks are part of the design. '
    'Your brain needs pauses to convert short-term memory to long-term memory.',
    label='8-HOUR RULE'
)

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 2 — YOUTUBE RESOURCES
# ══════════════════════════════════════════════════════════════════════════════
page_break()
heading1('Section 2 — YouTube Study Resources (Watch Alongside This Document)')
body(
    'Use these videos to reinforce what you read. Watch the relevant video BEFORE or DURING '
    'the theory block for that topic. All links below are free. '
    'Pause and re-watch sections you find unclear — that is not weakness, that is how engineers learn.',
    space_after=6
)

youtube_box([
    ('Linux Command Line — Complete Beginner Course (5 hrs)',  'freeCodeCamp',   'https://www.youtube.com/watch?v=ZtqBQ68cfJc'),
    ('Linux for Hackers — Full Series (Ep 1-10)',              'NetworkChuck',   'https://www.youtube.com/@NetworkChuck'),
    ('Linux File System Explained',                            'YouTube Search', 'https://www.youtube.com/results?search_query=linux+file+system+explained+for+beginners'),
    ('Linux File Permissions (chmod, chown) Deep Dive',        'YouTube Search', 'https://www.youtube.com/results?search_query=linux+file+permissions+chmod+chown+tutorial'),
    ('Linux Process Management (ps, kill, top)',               'YouTube Search', 'https://www.youtube.com/results?search_query=linux+process+management+ps+kill+top+tutorial'),
    ('vim Tutorial — Complete for Beginners',                  'YouTube Search', 'https://www.youtube.com/results?search_query=vim+tutorial+complete+beginners+2024'),
    ('Linux grep, awk, sed, pipe — Text Processing',           'YouTube Search', 'https://www.youtube.com/results?search_query=linux+grep+awk+sed+pipe+tutorial'),
    ('Two Sum — Brute Force to Optimal (NeetCode)',            'NeetCode',       'https://www.youtube.com/watch?v=KLlXCFG5TnA'),
    ('Big-O Notation Explained Simply',                        'YouTube Search', 'https://www.youtube.com/results?search_query=big+o+notation+explained+simply+beginners'),
])

tip_box(
    'For Linux basics, the freeCodeCamp video (first link) covers about 70% of today\'s content. '
    'For DSA, watch NeetCode\'s Two Sum explanation AFTER you attempt the problem yourself. '
    'NetworkChuck\'s Linux for Hackers series is great for motivation — he explains WHY, not just HOW.',
    label='RECOMMENDED WATCH ORDER'
)

heading2('Per-Section Video Guide')
simple_table(
    ['When You Study', 'Watch This'],
    [
        ['7:15–9:15 AM (Theory Block 1)', 'freeCodeCamp Linux full course — first 2 hours cover file system and commands'],
        ['9:30–11:00 AM (Theory Block 2)','Linux File Permissions link + Process Management link above'],
        ['During lunch (optional)',        'NetworkChuck Linux for Hackers Ep 1 — great motivation (23 min)'],
        ['12:00–1:00 PM (DSA block)',      'Attempt Two Sum first. THEN watch NeetCode video to check your approach.'],
        ['2:10–4:10 PM (Practical)',       'Keep freeCodeCamp video open as reference while running commands'],
    ]
)

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 3 — WHAT IS LINUX
# ══════════════════════════════════════════════════════════════════════════════
page_break()
heading1('Section 3 — What Is Linux and Why Every Server Runs It')

heading2('The Core Concept')
body(
    'Linux is an operating system kernel — the core software that sits between your hardware and '
    'your programs. It manages CPU scheduling, memory allocation, disk I/O, and network connections. '
    'Nothing runs without the kernel. Linux was written by Linus Torvalds in 1991 as a free, open-source '
    'alternative to Unix. Since the source code is public, companies like AWS built their entire '
    'infrastructure on it — saving billions in licensing costs.'
)

heading2('Why Servers Run Linux, Not Windows')
simple_table(
    ['Reason', 'What it means in practice'],
    [
        ['Zero cost',      'Windows Server license costs money per CPU core. Linux is free. At AWS scale = billions saved.'],
        ['Stability',      'Linux servers run for years without rebooting. Windows often requires restarts for updates.'],
        ['Security',       'No GUI by default. Fewer running processes = smaller attack surface = fewer exploits.'],
        ['Performance',    'Linux uses less RAM for the OS itself, leaving more RAM for your Spring Boot app.'],
        ['Scriptability',  '"Everything is a file" — you can automate and script every aspect of the system with bash.'],
        ['Container base', 'Docker, Kubernetes, and all container technology is built on Linux kernel features (cgroups, namespaces).'],
        ['Ecosystem',      'All DevOps tools — Ansible, Terraform, Jenkins, GitHub Actions runners — are Linux-native.'],
    ]
)

heading2('Linux Distributions You Will Actually Use')
simple_table(
    ['Distro', 'Package Manager', 'Where You See It', 'Base Image Size'],
    [
        ['Ubuntu 22.04 LTS',   'apt',       'Local dev, CI/CD runners, general servers',     '~75 MB'],
        ['Amazon Linux 2023',  'yum / dnf', 'AWS EC2 instances — you will use this heavily', '~165 MB'],
        ['CentOS / RHEL',      'yum',       'Enterprise servers, some corporate environments','~200 MB'],
        ['Alpine Linux',       'apk',       'Docker base images — keeps containers tiny',    '~5 MB'],
        ['Debian',             'apt',       'Stable servers, some cloud VMs',               '~120 MB'],
    ]
)

tip_box(
    'Interview answer for "which Linux have you used?": '
    '"Amazon Linux 2023 on EC2 for all our AWS workloads, and Ubuntu locally with Docker. '
    'I am comfortable with both apt and yum package managers."',
    label='INTERVIEW ANSWER'
)

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 4 — FILE SYSTEM
# ══════════════════════════════════════════════════════════════════════════════
heading1('Section 4 — The Linux File System Tree')

heading2('One Tree — No Drive Letters')
body(
    'Linux has a single tree starting at / (pronounced "root"). There are no C:\\ or D:\\ drives. '
    'Everything — hard disks, USB drives, network shares, virtual kernel interfaces — is mounted '
    'somewhere in this tree. This is confusing at first for Windows users. Accept it. It is '
    'actually more logical once you understand it.'
)

heading2('The Complete Directory Reference')
simple_table(
    ['Directory', 'Full Name / Purpose', 'Real-world use case'],
    [
        ['/',           'Root — top of the entire tree',                                         'Start all absolute paths here'],
        ['/home/',      'User home directories. cd ~ goes here.',                                'Your files, scripts, .bashrc'],
        ['/root/',      'Home dir of root (admin) user — NOT /home/root/',                      'Sudo operations, admin scripts'],
        ['/etc/',       'System-wide configuration files for ALL programs',                      'nginx.conf, hosts, environment, ssh/sshd_config'],
        ['/var/log/',   'Variable/log — constantly changing data. All system + app logs.',       'syslog, nginx/access.log, your app logs'],
        ['/var/lib/',   'Application state data — databases store files here',                   'MySQL data, Redis dump files'],
        ['/tmp/',       'Temporary files — cleared on EVERY reboot',                            'Quick scratch space during scripts'],
        ['/usr/bin/',   'User binary programs — most CLI tools live here',                       'java, python3, curl, git, grep'],
        ['/usr/local/', 'Locally compiled / manually installed software',                        'Tools you install from source'],
        ['/opt/',       'Optional/third-party software packages',                                'Datadog agent, custom vendor tools'],
        ['/proc/',      'Virtual filesystem — kernel exposes process info here. Not real files.','ps reads /proc, top reads /proc/meminfo'],
        ['/sys/',       'Virtual filesystem — hardware and kernel parameters',                   'Tuning kernel settings, driver info'],
        ['/dev/',       'Device files — /dev/sda = first hard disk, /dev/null = data void',     'Block devices, character devices'],
        ['/bin/',       'Essential system binaries needed before /usr mounts',                   'ls, cp, mv, bash, cat'],
        ['/sbin/',      'System admin binaries — most require root/sudo',                        'fdisk, ifconfig, init, reboot'],
        ['/mnt/',       'Mount point for temporary external drives / NFS shares',                'Mounting an EBS volume manually'],
        ['/boot/',      'Kernel image and bootloader files — do not touch',                      'GRUB, vmlinuz, initrd'],
    ]
)

heading2('The 6 Rules That Make Everything Click')
bullet('Rule 1: /etc = configuration. Program behaving wrong? Check /etc/<program-name>/ first.')
bullet('Rule 2: /var/log = logs. Something crashed? cd /var/log and grep for errors.')
bullet('Rule 3: /tmp = scratch space. Safe to write here, but do NOT store anything you need after reboot.')
bullet('Rule 4: ~ is always your home. cd ~ from anywhere brings you to /home/kanagarajan/.')
bullet('Rule 5: Everything is a file. Your hard disk is /dev/sda. Network sockets are file descriptors.')
bullet('Rule 6: /proc is virtual. Files in /proc do not exist on disk — the kernel generates them on demand.')

heading2('Absolute Paths vs Relative Paths')
body(
    'This distinction comes up in every script and every troubleshooting session. Always know which type you are using.'
)
simple_table(
    ['Type', 'Definition', 'Always starts with', 'Example', 'Works from where?'],
    [
        ['Absolute', 'Full path from root — completely unambiguous', '/', '/home/kanagarajan/projects/app.log', 'Anywhere on the system'],
        ['Relative', 'Path relative to your CURRENT directory', 'Anything other than /', './app.log  or  ../configs/', 'Only correct from specific location'],
    ]
)
code_block([
    '# You are currently in: /home/kanagarajan/projects/',
    '',
    '# Absolute paths — always work regardless of where you are:',
    'cat /home/kanagarajan/projects/myapp/app.log',
    'cd /etc/nginx/',
    '',
    '# Relative paths — depend on current directory:',
    'cat ./myapp/app.log        # ./ = current directory (the ./ is optional)',
    'cat myapp/app.log          # same result — ./ is implied when omitted',
    'cd ..                      # go UP one level to /home/kanagarajan/',
    'cd ../..                   # go UP two levels to /home/',
    '',
    '# Special shortcuts:',
    '.     = current directory',
    '..    = parent directory',
    '~     = your home (/home/kanagarajan/)',
    '-     = previous directory (cd - toggles between last two locations)',
])
tip_box(
    'In shell scripts, ALWAYS use absolute paths. Scripts can be called from anywhere, '
    'so a relative path like ./app.log could point to different files depending on '
    'where the script was called from. Bugs from this mistake are hard to find.',
    label='SCRIPTING RULE'
)

qa_block(
    'What is the difference between an absolute path and a relative path in Linux?',
    'An absolute path starts from the root directory (/) and gives the complete, unambiguous '
    'location of a file — for example /etc/nginx/nginx.conf. It works from any directory on the system. '
    'A relative path is relative to your current working directory — for example ./config.properties '
    'or ../logs/app.log. It only works correctly if you are in the right directory. '
    'In shell scripts I always use absolute paths because the script may be invoked from any location '
    '(cron jobs, systemd, CI/CD pipelines) and a relative path would point to the wrong place. '
    'Interactively I use relative paths for faster navigation.'
)

qa_block(
    'What is the purpose of the /etc directory in Linux?',
    '/etc contains system-wide configuration files for all installed programs. For example, /etc/nginx/nginx.conf '
    'configures Nginx, /etc/hosts overrides DNS lookups locally, /etc/environment sets system-wide environment '
    'variables, and /etc/ssh/sshd_config controls SSH server settings. When a service is misbehaving on a server, '
    'the first place I look is /etc/<service-name>/. It is also common to back up /etc/ before making system changes.'
)

qa_block(
    'What is the /proc filesystem and why is it special?',
    '/proc is a virtual filesystem — it does not correspond to real files stored on disk. The kernel generates '
    'its contents on the fly to expose information about running processes and kernel state. /proc/1234/ is a '
    'directory for process 1234 containing its memory maps, open file descriptors, and environment variables. '
    '/proc/cpuinfo has CPU details, /proc/meminfo has memory stats. Tools like top, ps, and free all read '
    'from /proc. This is the Linux "everything is a file" philosophy — even kernel internals are accessible '
    'as regular file reads, making monitoring and scripting straightforward.'
)

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 5 — ESSENTIAL COMMANDS
# ══════════════════════════════════════════════════════════════════════════════
page_break()
heading1('Section 5 — Essential Commands (Learn These Cold)')

heading2('Navigation Commands')
code_block([
    'pwd                        # Print Working Directory — where am I right now?',
    'ls                         # List files and folders in current directory',
    'ls -l                      # Long format: permissions, owner, size, modified date',
    'ls -la                     # Long format + hidden files (files starting with .)',
    'ls -lh                     # Long format + human-readable sizes (4.0K, 2.5M, 1.2G)',
    'ls -lt                     # Long format + sorted by modification time (newest first)',
    'ls -lS                     # Long format + sorted by file size (largest first)',
    'cd /path/to/dir            # Change into a specific directory using absolute path',
    'cd ..                      # Go up one level (parent directory)',
    'cd ~                       # Go to your home directory',
    'cd -                       # Toggle: go back to the previous directory',
])

heading3('Anatomy of ls -l Output')
code_block([
    '-rwxr-xr-x  1  ec2-user  ec2-user  4096  May 15 09:30  deploy.sh',
    ' [1]        [2]  [3]      [4]       [5]   [6]           [7]',
    '',
    '[1] -rwxr-xr-x  = File type + permissions (- regular file, d directory, l symlink)',
    '[2] 1            = Number of hard links to this file',
    '[3] ec2-user     = Owner (the user who owns the file)',
    '[4] ec2-user     = Group owner',
    '[5] 4096         = File size in bytes (use ls -lh to see KB/MB/GB)',
    '[6] May 15 09:30 = Last modification date and time',
    '[7] deploy.sh    = File name',
])

heading2('File and Directory Operations')
code_block([
    'mkdir mydir                # Create a directory',
    'mkdir -p a/b/c             # Create nested directories — -p creates parents if missing',
    'touch myfile.txt           # Create empty file OR update its timestamp if it exists',
    'cp source.txt dest.txt     # Copy a file',
    'cp -r sourcedir/ destdir/  # Copy a directory and ALL its contents recursively',
    'cp -p source.txt dest.txt  # Copy and PRESERVE timestamps and permissions',
    'mv oldname.txt newname.txt # Rename a file (same directory)',
    'mv file.txt /tmp/          # Move a file to another directory',
    'rm file.txt                # Delete a file — PERMANENT. No recycle bin.',
    'rm -r mydir/               # Delete a directory and everything inside it',
    'rm -rf mydir/              # Force delete without confirmation — DANGEROUS',
    'rm -i file.txt             # Interactive: ask before each deletion (safer)',
])

important_box(
    'rm -rf with a wrong path can destroy your OS. ALWAYS triple-check before running. '
    'Never run: rm -rf /  or  rm -rf /*  — these delete everything on the server.',
    label='DANGER'
)

heading2('Symbolic Links and Hard Links')
body(
    'A symbolic link (symlink) is a pointer \u2014 a special file that contains the path to another file or directory. '
    'A hard link is a second directory entry pointing to the EXACT SAME data on disk (same inode).'
)
simple_table(
    ['', 'Hard Link', 'Symbolic Link (Symlink)'],
    [
        ['Create with',          'ln target linkname',               'ln -s target linkname'],
        ['Points to',            'Same inode \u2014 identical disk data',  'Stores the PATH of the target'],
        ['If original deleted',  'Data still accessible \u2014 safe',      'Link breaks (dangling link)'],
        ['Cross filesystem?',    'No \u2014 same partition only',           'Yes \u2014 can cross disks'],
        ['ls -la shows',         'Same as a regular file',            'lrwxrwxrwx with \u2192 path'],
        ['Common use',           'Backup / hard delete protection',   'PATH aliases, version switching'],
    ]
)
code_block([
    '# Create a symbolic link:',
    'ln -s /opt/java-21/bin/java /usr/local/bin/java',
    '    # Now typing "java" runs /opt/java-21/bin/java',
    '',
    '# ls -la shows symlinks with l at position 1 and \u2192 showing the target:',
    'lrwxrwxrwx  1 root root  24 May 15  java -> /opt/java-21/bin/java',
    '',
    '# Check where a symlink points:',
    'readlink /usr/local/bin/java         # Shows: /opt/java-21/bin/java',
    'readlink -f /usr/local/bin/java      # Fully resolved absolute path',
    '',
    '# Remove a symlink (use rm, NOT rmdir \u2014 it is a file, not a real directory):',
    'rm /usr/local/bin/java               # Removes the LINK, not the target file',
    '',
    '# Create a hard link:',
    'ln existing_file.txt hardlink.txt    # Both names now point to the same data',
    '# Deleting existing_file.txt does NOT delete the data \u2014 hardlink.txt still works',
])
qa_block(
    'What is the difference between a hard link and a symbolic link in Linux?',
    'A hard link (ln source dest) creates a second directory entry that points to the SAME inode \u2014 '
    'the same physical data blocks on disk. Deleting the original file does not affect the hard link '
    'because the data persists until every hard link referencing that inode is removed. '
    'Hard links cannot cross filesystem boundaries and cannot point to directories. '
    'A symbolic link (ln -s source dest) is a pointer that stores the PATH of the target file. '
    'If the target is deleted or moved, the symlink becomes a dangling broken link. '
    'Symlinks can cross filesystems and can point to directories. '
    'In practice I use symlinks to create version aliases \u2014 for example when switching Java versions: '
    'ln -s /opt/java-21/bin/java /usr/local/bin/java. Updating just the symlink switches the version '
    'for all users without changing any application config.'
)

heading2('Reading File Contents')
code_block([
    'cat file.txt               # Print ENTIRE file to terminal (use for small files)',
    'cat -n file.txt            # Print file with line numbers',
    'less file.txt              # Read page by page: q=quit  /text=search  n=next match  G=end',
    'head file.txt              # Show first 10 lines',
    'head -n 20 file.txt        # Show first 20 lines',
    'tail file.txt              # Show last 10 lines',
    'tail -n 50 file.txt        # Show last 50 lines',
    'tail -f app.log            # FOLLOW a log in real time — new lines appear as written',
    '                           # Press Ctrl+C to stop following',
    'tail -f app.log | grep ERROR  # Follow log but only show ERROR lines live',
])

tip_box(
    '"tail -f /var/log/app.log" is how every DevOps engineer monitors a running application. '
    'Combine with grep to filter noise: tail -f app.log | grep -i error',
    label='PRODUCTION TIP'
)

heading2('Searching — grep and find')
code_block([
    '# ── grep: search INSIDE files ─────────────────────────────────────────',
    'grep "ERROR" app.log               # Lines containing ERROR (case-sensitive)',
    'grep -i "error" app.log            # Case-insensitive search',
    'grep -n "NullPointer" app.log      # Show line NUMBER with each match',
    'grep -c "ERROR" app.log            # Count of matching lines only (a number)',
    'grep -v "DEBUG" app.log            # Lines that do NOT contain DEBUG',
    'grep -r "datasource" /etc/         # Recursive: search inside all files in /etc/',
    'grep -l "ERROR" *.log             # List only FILE NAMES that contain ERROR',
    'grep -E "ERROR|WARN" app.log       # Regex: match ERROR or WARN',
    'grep -A 3 "Exception" app.log      # Show 3 lines AFTER each match (context)',
    'grep -B 2 "Exception" app.log      # Show 2 lines BEFORE each match',
    'grep -o "ERROR\\|WARN" app.log      # Output ONLY the matching text, not the whole line',
    '                                    # Useful for extracting specific patterns from logs',
    '',
    '# ── find: search for FILES by name/type/date ───────────────────────────',
    'find / -name "*.java"              # Find all .java files (from root)',
    'find . -name "*.log"               # Find .log files in current directory',
    'find /home -name "*.log" -mtime -7 # .log files modified in last 7 days',
    'find /tmp -type d                  # Find only DIRECTORIES inside /tmp',
    'find /tmp -type f -size +10M       # Find files larger than 10 MB',
    'find . -name "*.conf" -exec cat {} \\;  # Find .conf files and print their contents',
])

heading2('Creating and Writing Files')
code_block([
    'echo "Hello World"                 # Print text to terminal (stdout)',
    'echo "Hello World" > file.txt      # Write to file — OVERWRITES entire file',
    'echo "Second line" >> file.txt     # Append to file — adds to end, SAFE',
    'cat > file.txt                     # Type content directly, Ctrl+D to save',
    'cat >> file.txt                    # Append by typing, Ctrl+D to save',
])

important_box(
    '> OVERWRITES the file completely. >> APPENDS to the end. '
    'In DevOps, always ask yourself: "Do I want to replace or add?" before using >.',
    label='CRITICAL DIFFERENCE'
)

heading2('System Information — Commonly Asked in Interviews')
code_block([
    '# ── Disk usage ─────────────────────────────────────────────────────────',
    'df -h                       # Disk space of all mounted filesystems (human readable)',
    'df -h /                     # Disk space of the root filesystem specifically',
    'du -sh *                    # Size of each file/folder in current directory',
    'du -sh /var/log/            # Total size of /var/log/ directory',
    'du -sh * | sort -rh | head -10  # Top 10 largest items in current directory',
    '',
    '# ── Memory usage ───────────────────────────────────────────────────────',
    'free -h                     # RAM usage: total, used, free, shared, buff/cache',
    '                            # "available" column = RAM your app can actually use',
    '',
    '# ── CPU and system ─────────────────────────────────────────────────────',
    'uname -a                    # Kernel version and system architecture',
    'hostname                    # Server hostname',
    'uptime                      # How long server has been running + load averages',
    'whoami                      # Current user',
    'id                          # Current user + all their group memberships',
    'which java                  # Find the full path of where java is installed',
    '',
    '# ── Command history ────────────────────────────────────────────────────',
    'history                     # Show last 1000 commands you ran',
    'history | grep grep         # Find grep commands you ran before',
    '!!                          # Re-run the LAST command',
    'Ctrl + R                    # Interactive reverse search through command history',
    '',
    '# ── Getting help ───────────────────────────────────────────────────────',
    'man ls                      # Full manual page for the ls command (q to quit)',
    'ls --help                   # Quick help summary (shorter than man)',
])

tip_box(
    'df -h and free -h are your first commands when a server is "slow" or "out of memory". '
    'du -sh * helps you find what is eating disk space. Mention these in interviews when asked '
    'about troubleshooting server performance.',
    label='TROUBLESHOOTING COMMANDS'
)

qa_block(
    'How do you check if a server is running out of disk space?',
    'I run df -h to see disk usage across all filesystems — it shows total, used, and available space '
    'in human-readable format. If I see a filesystem at 90%+ usage, I run du -sh * in the full '
    'directories to find what is consuming space — usually /var/log/ growing from unrotated logs, '
    'or /tmp/ from uncleaned temporary files. I also check free -h to see memory usage. '
    'On production servers I set up CloudWatch disk and memory alarms to alert before we hit 80%.'
)

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 6 — FILE PERMISSIONS
# ══════════════════════════════════════════════════════════════════════════════
page_break()
heading1('Section 6 — File Permissions (Interviewers Test This Heavily)')

youtube_box([
    ('Linux File Permissions Explained (chmod, chown)', 'YouTube Search', 'https://www.youtube.com/results?search_query=linux+file+permissions+chmod+chown+explained+tutorial'),
    ('Linux chmod Numeric Mode — Full Guide',           'YouTube Search', 'https://www.youtube.com/results?search_query=linux+chmod+numeric+mode+755+644+tutorial'),
])

heading2('Three Categories, Three Permission Types')
body('Every file and directory has permissions for three categories of users:')
bullet('Owner (u) — the user who owns the file')
bullet('Group (g) — a group of users who share access')
bullet('Others (o) — everyone else on the system')

simple_table(
    ['Permission', 'Symbol', 'Numeric', 'On a FILE means', 'On a DIRECTORY means'],
    [
        ['Read',    'r', '4', 'View file contents with cat/less',   'List files inside (ls)'],
        ['Write',   'w', '2', 'Modify the file content',            'Create or delete files inside'],
        ['Execute', 'x', '1', 'Run the file as a program',          'Enter the directory with cd'],
        ['None',    '-', '0', 'No access',                          'No access'],
    ]
)

heading2('Reading the Permission String — Full Breakdown')
code_block([
    'Example: -rwxr-xr--',
    '',
    'Position 1:    -    = File type:  - regular file  d directory  l symbolic link',
    'Positions 2-4: rwx  = Owner permissions: r(4) + w(2) + x(1) = 7',
    'Positions 5-7: r-x  = Group permissions: r(4) + -(0) + x(1) = 5',
    'Positions 8-10: r-- = Others permissions: r(4) + -(0) + -(0) = 4',
    '',
    'So -rwxr-xr-- in numeric notation = 754',
    '',
    'drwxr-xr-x  = directory (d), owner=7, group=5, others=5 → numeric: 755',
    'lrwxrwxrwx  = symbolic link (l), usually 777 — link target has real permissions',
])

heading2('chmod — Changing Permissions')
code_block([
    '# ── Numeric notation (most common, fastest) ────────────────────────────',
    'chmod 755 deploy.sh        # Owner: rwx(7), Group: r-x(5), Others: r-x(5)',
    'chmod 644 config.txt       # Owner: rw-(6), Group: r--(4), Others: r--(4)',
    'chmod 600 private.pem      # Owner: rw-(6), Group: ---(0), Others: ---(0)',
    'chmod 700 my_script.sh     # Owner: rwx(7), Group: ---(0), Others: ---(0)',
    'chmod 777 shared/          # Everyone: rwx — AVOID on production (security risk)',
    '',
    '# ── Symbolic notation (readable, good for targeted changes) ────────────',
    'chmod +x script.sh         # Add execute for ALL (owner, group, others)',
    'chmod -x script.sh         # Remove execute for ALL',
    'chmod u+x script.sh        # Add execute for OWNER only',
    'chmod g-w file.txt         # Remove write from GROUP only',
    'chmod o-r secret.txt       # Remove read from OTHERS only',
    '',
    '# ── Recursive ──────────────────────────────────────────────────────────',
    'chmod -R 755 /app/scripts/ # Set 755 on all files in /app/scripts/',
])

heading2('Common Permission Patterns and Why')
simple_table(
    ['Permissions', 'Numeric', 'Use Case', 'Why'],
    [
        ['rwx------', '700', 'Private personal script',       'Only you can read/write/run it'],
        ['rw-------', '600', 'SSH private key, .env secrets', 'SSH refuses keys readable by others'],
        ['rwxr-xr-x', '755', 'Public scripts, binaries',      'Anyone can run, only owner edits'],
        ['rw-r--r--', '644', 'Public config, web files',      'Anyone reads, only owner writes'],
        ['rw-rw-r--', '664', 'Shared group files',            'Owner and group write, others read'],
        ['rwxrwxrwx', '777', 'Avoid on production',           'Anyone can modify — security hole'],
    ]
)

heading2('chown and sudo')
code_block([
    'chown kanagarajan file.txt                # Change owner to kanagarajan',
    'chown kanagarajan:developers file.txt     # Change owner AND group in one command',
    'chown -R kanagarajan:developers /app/     # Recursively change for all files in /app/',
    'sudo chown www-data:www-data /var/www/html/  # Give nginx ownership of web files',
    '',
    'sudo command               # Run a single command as root (asks for YOUR password)',
    'sudo apt install nginx     # Install a package (requires root)',
    'sudo !!                    # Re-run last command with sudo (when you forgot)',
])

qa_block(
    'What does chmod 755 mean? Explain every digit.',
    'chmod 755 sets the file permission to rwxr-xr-x. Each digit represents permissions for '
    'one category of user. The first digit 7 is for the OWNER: read(4) + write(2) + execute(1) = 7, '
    'so the owner has full access — they can read, modify, and run the file. The second digit 5 is '
    'for the GROUP: read(4) + execute(1) = 5 — group members can read and run but cannot modify. '
    'The third digit 5 is for OTHERS (everyone else): same as group. '
    'chmod 755 is the standard permission for executable scripts and public binaries.'
)

qa_block(
    'Why must an SSH private key be chmod 600?',
    'The SSH client explicitly checks the permissions of your private key file before using it. '
    'If the key is readable by anyone other than the owner (e.g., 644 or 755), SSH prints '
    '"WARNING: UNPROTECTED PRIVATE KEY FILE!" and refuses to connect. This is a deliberate '
    'security measure — if another user on the same server can read your private key, they can '
    'impersonate you on remote servers. chmod 600 (rw-------) means ONLY the owner can read or '
    'write the file, satisfying SSH\'s security requirement.'
)

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 7 — PROCESS MANAGEMENT
# ══════════════════════════════════════════════════════════════════════════════
page_break()
heading1('Section 7 — Process Management')

youtube_box([
    ('Linux Process Management — ps, kill, top, htop', 'YouTube Search', 'https://www.youtube.com/results?search_query=linux+process+management+ps+kill+top+htop+tutorial'),
    ('Linux Signals — SIGTERM vs SIGKILL explained',    'YouTube Search', 'https://www.youtube.com/results?search_query=linux+signals+sigterm+sigkill+explained'),
])

heading2('Viewing Processes')
code_block([
    'ps                         # Show processes owned by current user in THIS terminal',
    'ps aux                     # Show ALL processes from ALL users (most useful)',
    '                           # a=all users  u=user-oriented format  x=background processes',
    'ps aux | grep java         # Find all Java processes',
    'ps aux | grep -v grep      # Remove the grep process itself from results',
    'ps -ef                     # Another format: shows PPID (parent PID) as well',
    'ps aux --sort=-%cpu | head # Top 10 processes by CPU usage',
    'ps aux --sort=-%mem | head # Top 10 processes by memory usage',
    '',
    'top                        # Real-time interactive process viewer (updates every 3 sec)',
    '                           # Keys: q=quit  k=kill  M=sort by memory  P=sort by CPU',
    'htop                       # Better interactive viewer (sudo yum install htop)',
])

heading2('Understanding ps aux Output')
code_block([
    'USER    PID  %CPU  %MEM     VSZ    RSS  TTY   STAT  START   TIME  COMMAND',
    'ec2-user 1234 25.0  8.5 2048000 174000  pts/0  Sl   09:15   2:30  java -jar app.jar',
    '',
    'PID    = Process ID. The unique number to kill or signal the process.',
    '%CPU   = Percentage of CPU time used right now.',
    '%MEM   = Percentage of total RAM used right now.',
    'RSS    = Resident Set Size — ACTUAL physical RAM in use. Use this for real memory.',
    'STAT codes:',
    '  S = Sleeping (waiting for I/O or event)',
    '  R = Running (using CPU right now)',
    '  Z = Zombie (process finished but parent has not cleaned it up)',
    '  D = Uninterruptible sleep (waiting for disk I/O — cannot be killed)',
    '  T = Stopped (paused, e.g., by Ctrl+Z)',
])

heading2('Killing and Signaling Processes')
code_block([
    'kill 1234                  # Send SIGTERM (15) to PID 1234 — polite graceful shutdown',
    'kill -9 1234               # Send SIGKILL (9) — immediate force kill, no cleanup',
    'kill -HUP 1234             # SIGHUP — tells many daemons to reload config without restart',
    'killall java               # Kill ALL processes named "java" — use carefully',
    'pkill -f "app.jar"         # Kill all processes whose full command matches "app.jar"',
    '',
    '# Signals reference:',
    '# SIGHUP  (1)  = Hangup — reload config (nginx, logrotate use this)',
    '# SIGTERM (15) = Graceful shutdown request — catchable by the application',
    '# SIGKILL (9)  = Immediate kill — cannot be caught or ignored by the process',
])

important_box(
    'ALWAYS try kill <PID> (SIGTERM) first. Give it 30 seconds. '
    'Only escalate to kill -9 if the process does not respond. '
    'SIGKILL leaves DB connections leaked, in-flight requests dropped, and files unclosed.',
    label='SIGTERM BEFORE SIGKILL'
)

heading2('Background Processes — nohup')
code_block([
    'java -jar app.jar &            # & sends process to background — terminal stays free',
    'nohup java -jar app.jar &      # nohup prevents SIGHUP when you close SSH session',
    '                               # Without nohup: closing SSH kills the process',
    'nohup java -jar app.jar > app.log 2>&1 &   # Redirect ALL output to app.log',
    '                                           # 2>&1 = merge stderr into stdout',
    '',
    'jobs                           # List background jobs in current terminal',
    'fg %1                          # Bring job 1 back to FOREGROUND (resume interactively)',
    'bg %1                          # Resume a Ctrl+Z stopped job IN THE BACKGROUND',
    '',
    '# nohup.out \u2014 what happens when you do not specify a log file:',
    'nohup java -jar app.jar &',
    '    # Output goes to ./nohup.out in the current directory by default',
    '    # Always redirect explicitly to avoid confusion:',
    'nohup java -jar app.jar > app.log 2>&1 &',
    '    # Now output + errors go to app.log, not nohup.out',
    'cat nohup.out                  # Check output if you forgot to redirect',
    '',
    'Ctrl + Z                       # PAUSE (stop) the current foreground process',
    'Ctrl + C                       # TERMINATE the current foreground process (SIGINT)',
])

heading2('Port and Socket Investigation')
code_block([
    'lsof -i :8080              # Which process is using port 8080?',
    'lsof -i :3306              # Which process is using MySQL port 3306?',
    'ss -tlnp                   # Modern replacement for netstat — list listening TCP ports',
    'ss -tlnp | grep 8080       # Check who is listening on port 8080',
    'netstat -tlnp              # Older equivalent to ss',
])

tip_box(
    '"Port 8080 is already in use" when starting Spring Boot? Run: lsof -i :8080 '
    'Get the PID from the output, then: kill <PID>. Then retry starting the app.',
    label='SPRING BOOT TIP'
)

qa_block(
    'How do you find a process and kill it gracefully in Linux?',
    'I use ps aux | grep java to find the process — this shows all running processes filtered '
    'to ones with "java" in their command line. The PID is in the second column. '
    'I use kill <PID> to send SIGTERM, which is a graceful shutdown request. The application '
    'receives it and can finish in-flight requests, close database connections, and run any '
    '@PreDestroy cleanup methods in Spring Boot. I wait 30 seconds. If it still runs, I escalate '
    'to kill -9 <PID> to force-kill it immediately, accepting that open resources may be left '
    'in an inconsistent state. For checking port conflicts, I use lsof -i :8080.'
)

qa_block(
    'What is the difference between SIGTERM and SIGKILL?',
    'SIGTERM (signal 15) is a polite shutdown request that the application can catch and handle. '
    'When received, a well-written application like Spring Boot will stop accepting new requests, '
    'complete in-progress work, flush buffers, close DB connections, and run cleanup code '
    'such as @PreDestroy methods. SIGKILL (signal 9) is fundamentally different — it cannot '
    'be caught, blocked, or ignored by the application. The kernel terminates the process '
    'immediately with no cleanup. This means database connection pools may leak, files can '
    'be left in a partially-written corrupt state, and any in-flight HTTP requests are '
    'dropped abruptly. I always use SIGTERM first and only escalate to SIGKILL as a last resort.'
)

qa_block(
    'What is nohup and when would you use it?',
    'nohup stands for "no hangup". By default, when you close an SSH session, the shell sends '
    'a SIGHUP signal to all processes started in that session, and they terminate. '
    'nohup java -jar app.jar & prevents the process from receiving SIGHUP, so it keeps running '
    'after you disconnect. The & sends it to the background so the terminal prompt returns. '
    'I prefer to redirect output explicitly: nohup java -jar app.jar > app.log 2>&1 & '
    'In production I use systemctl to manage services properly — it handles restart-on-failure, '
    'logging, and boot-time startup. But nohup is useful for quick ad-hoc long-running jobs.'
)

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 8 — TEXT EDITORS
# ══════════════════════════════════════════════════════════════════════════════
page_break()
heading1('Section 8 — Text Editors: nano and vim')

youtube_box([
    ('vim Tutorial — Complete Beginners Guide 2024', 'YouTube Search', 'https://www.youtube.com/results?search_query=vim+tutorial+complete+beginners+2024'),
    ('nano Editor Basics — Quick Start',             'YouTube Search', 'https://www.youtube.com/results?search_query=nano+editor+linux+tutorial+beginners'),
])

body(
    'On a production server there is no VS Code, no graphical interface — only the terminal. '
    'You must be able to edit a config file quickly when a service is down and seconds matter. '
    'Learn both: nano to get work done safely, vim because it is on EVERY server.',
    space_after=6
)

heading2('nano — The Safe Choice')
body('nano is straightforward — what you type appears directly. No mode-switching confusion.')
code_block([
    'nano filename.txt          # Open or create a file',
    'sudo nano /etc/hosts       # Edit a system file requiring root',
])

simple_table(
    ['Key Combination', 'Action'],
    [
        ['Type normally',       'Text appears where cursor is. No special mode.'],
        ['Ctrl + O, then Enter', 'Save the file (Write Out)'],
        ['Ctrl + X',            'Exit — asks to save if unsaved changes'],
        ['Ctrl + K',            'Cut (delete) the current line'],
        ['Ctrl + U',            'Paste the cut line'],
        ['Ctrl + W',            'Search (Where is?) — type search term, Enter'],
        ['Alt + U',             'Undo last action'],
    ]
)

heading2('vim — Modes Are the Key Concept')
body(
    'Vim is installed on virtually every Linux server in existence. '
    'The key insight: it has MODES. You cannot type directly when vim opens.'
)

simple_table(
    ['Mode', 'How to Enter', 'What You Can Do'],
    [
        ['Normal Mode',  'Default when vim opens. Press Esc anytime to return.', 'Navigate, delete, copy — every key is a command'],
        ['Insert Mode',  'Press i (before cursor), a (after cursor), o (new line below)', 'Type text normally like any editor'],
        ['Visual Mode',  'Press v (character), V (line)',                         'Select text for copy/cut/indent'],
        ['Command Mode', 'Press : from Normal Mode',                              'Save, quit, search, replace'],
    ]
)

heading3('Complete vim Survival Kit')
code_block([
    '# ── Opening and switching modes ────────────────────────────────────────',
    'vim filename.txt           # Open file',
    'i                          # Enter Insert Mode BEFORE cursor',
    'a                          # Enter Insert Mode AFTER cursor',
    'o                          # Open NEW LINE below and enter Insert Mode',
    'Esc                        # Return to Normal Mode (press this when lost)',
    '',
    '# ── Saving and quitting (from Normal Mode, type :) ─────────────────────',
    ':w                         # Save (write) the file',
    ':q                         # Quit — only works if no unsaved changes',
    ':wq                        # Save AND quit',
    ':q!                        # Quit WITHOUT saving (discard all changes)',
    '',
    '# ── Navigation in Normal Mode ───────────────────────────────────────────',
    'gg                         # Jump to FIRST line of file',
    'G                          # Jump to LAST line of file',
    ':42                        # Jump to line 42',
    '0                          # Jump to start of current line',
    '$                          # Jump to end of current line',
    '',
    '# ── Editing in Normal Mode ──────────────────────────────────────────────',
    'dd                         # Delete (cut) current line',
    'yy                         # Yank (copy) current line',
    'p                          # Paste below current line',
    'u                          # Undo',
    'Ctrl + R                   # Redo',
    '',
    '# ── Search and replace ──────────────────────────────────────────────────',
    '/search_word               # Search forward for search_word (n=next, N=previous)',
    ':%s/old/new/g              # Replace ALL occurrences of old with new in file',
    ':%s/old/new/gc             # Replace with confirmation for each one',
])

tip_box(
    'Minimum to remember: open vim, press i, type, press Esc, type :wq, press Enter. '
    'That is enough for 80% of real server situations. The rest comes with practice. '
    'Watch the YouTube search link for "vim tutorial beginners 2024" to see it in action.',
    label='VIM MINIMUM'
)

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 9 — PIPING AND TEXT PROCESSING
# ══════════════════════════════════════════════════════════════════════════════
page_break()
heading1('Section 9 — Piping, Redirection, and Text Processing')

youtube_box([
    ('Linux Pipe and Redirection — Complete Tutorial',  'YouTube Search', 'https://www.youtube.com/results?search_query=linux+pipe+redirection+tutorial+stdout+stderr'),
    ('grep awk sed — Linux Text Processing Full Guide', 'YouTube Search', 'https://www.youtube.com/results?search_query=linux+grep+awk+sed+text+processing+tutorial'),
])

heading2('stdin, stdout, stderr — The Three Standard Streams')
body(
    'Every Linux process has three standard data streams. Understanding these is essential to '
    'understanding how pipes, redirection, and 2>&1 actually work.'
)
simple_table(
    ['Stream', 'File Descriptor', 'What It Is', 'Default Destination'],
    [
        ['stdin  (standard input)',  '0', 'Where a process reads INPUT from',  'Your keyboard'],
        ['stdout (standard output)', '1', 'Where normal OUTPUT goes',          'Your terminal screen'],
        ['stderr (standard error)',  '2', 'Where ERROR messages and diagnostics go', 'Your terminal screen (separate from stdout)'],
    ]
)
code_block([
    '# File descriptors: stdin=0, stdout=1, stderr=2',
    '',
    '# Redirection operators work with these file descriptors:',
    'cmd > out.txt              # Redirect stdout (1) to a file -- overwrites',
    'cmd 1> out.txt             # Same thing -- 1 is the file descriptor for stdout',
    'cmd 2> err.txt             # Redirect stderr (2) to a SEPARATE file',
    'cmd > out.txt 2>&1         # Redirect stdout to file, THEN redirect stderr',
    '                           # to wherever stdout currently points (= out.txt)',
    '                           # Result: both normal output AND errors go to out.txt',
    'cmd 2>/dev/null            # Discard all error messages -- silence stderr',
    '',
    '# The pipe | connects stdout of left command to stdin of right command:',
    'cat app.log | grep "ERROR" # cat writes to stdout -> pipe -> grep reads from stdin',
    '                           # ONLY stdout is piped. stderr still goes to screen.',
    'cmd 2>&1 | grep "text"     # Merge stderr into stdout FIRST, then pipe both',
])
important_box(
    '"cmd > out.txt 2>&1" -- ORDER MATTERS. First redirect stdout to out.txt. '
    'THEN 2>&1 redirects stderr to where stdout now points (the file). '
    'If you reverse: "cmd 2>&1 > out.txt" -- stderr redirects to where stdout was '
    'at THAT moment (the terminal), THEN stdout goes to the file. '
    'Stderr would still appear on screen. Always write 2>&1 AFTER the > redirection.',
    label='2>&1 ORDER MATTERS'
)

qa_block(
    'What does 2>&1 mean? How does it work with > redirection?',
    'Linux gives every process three standard streams: stdin (fd 0), stdout (fd 1), and stderr (fd 2). '
    '2>&1 means redirect file descriptor 2 (stderr) to wherever file descriptor 1 (stdout) currently points. '
    'In "nohup java -jar app.jar > app.log 2>&1 &": '
    'first > app.log redirects stdout to the file. '
    'Then 2>&1 redirects stderr to wherever stdout now points, which is app.log. '
    'Result: both normal output and stack traces are captured in app.log. '
    'Without 2>&1, error messages would still appear on screen even though normal output goes to the file. '
    'The order matters -- 2>&1 must come AFTER the > redirection, not before it.'
)

heading2('The Pipe — | — The Most Powerful Concept in Linux')
body(
    'The pipe | takes the OUTPUT of the left command and feeds it as INPUT to the right command. '
    'You can chain as many pipes as you want. This lets you build complex data transformations '
    'out of simple, single-purpose tools — the Unix philosophy.'
)
code_block([
    'cat app.log | grep "ERROR"                      # Filter log to ERROR lines only',
    'cat app.log | grep "ERROR" | wc -l              # Count ERROR lines',
    'cat app.log | grep "ERROR" | sort | uniq -c     # Count unique error messages',
    'ps aux | grep java | grep -v grep               # Find Java processes (remove grep itself)',
    'ls -la | sort -k5 -n                            # List files sorted by size (column 5)',
    'history | grep "kubectl"                        # Find kubectl commands from your history',
])

heading2('Redirection — Where Output Goes')
simple_table(
    ['Operator', 'Meaning', 'Example', 'What Happens'],
    [
        ['>',        'Redirect stdout — OVERWRITE file',    'cmd > out.txt',       'Creates/overwrites out.txt'],
        ['>>',       'Redirect stdout — APPEND to file',    'cmd >> out.txt',      'Adds to end of out.txt'],
        ['2>',       'Redirect stderr (errors) to file',    'cmd 2> err.log',      'Errors go to err.log'],
        ['2>&1',     'Redirect stderr INTO stdout stream',  'cmd > all.log 2>&1',  'Both output+errors → all.log'],
        ['/dev/null', 'Discard output completely',          'cmd > /dev/null',     'Output silently discarded'],
    ]
)

heading2('Text Processing Commands')
code_block([
    '# ── wc (word count) ────────────────────────────────────────────────────',
    'wc -l file.txt              # Count LINES',
    'wc -w file.txt              # Count WORDS',
    'wc -c file.txt              # Count BYTES',
    '',
    '# ── sort ───────────────────────────────────────────────────────────────',
    'sort file.txt               # Sort lines alphabetically (A-Z)',
    'sort -r file.txt            # Reverse sort (Z-A)',
    'sort -n numbers.txt         # Numeric sort (not alphabetical)',
    'sort -rn numbers.txt        # Numeric reverse (largest first)',
    'sort -u file.txt            # Sort AND remove duplicates in one step',
    '',
    '# ── uniq ───────────────────────────────────────────────────────────────',
    'sort file.txt | uniq        # Sort then remove all duplicates',
    'sort file.txt | uniq -c     # Count occurrences of each unique line',
    'sort file.txt | uniq -c | sort -rn  # Rank by frequency (most common first)',
    '',
    '# ── cut ────────────────────────────────────────────────────────────────',
    'cut -d: -f1 /etc/passwd     # Get first field, colon-delimited (usernames)',
    'cut -d, -f2 data.csv        # Get second column from CSV',
    '',
    '# ── awk ────────────────────────────────────────────────────────────────',
    'awk \'{print $1}\' file.txt   # Print first space-delimited column',
    'awk \'{print $1, $3}\' f.txt  # Print columns 1 and 3',
    'awk -F: \'{print $1}\' /etc/passwd  # Use : as delimiter, print field 1',
    '',
    '# ── sed ────────────────────────────────────────────────────────────────',
    'sed "s/old/new/g" file.txt        # Replace ALL occurrences (print only)',
    'sed -i "s/old/new/g" file.txt     # Replace in-place (EDITS THE ACTUAL FILE)',
    'sed -i.bak "s/old/new/g" file.txt # Replace and save backup as file.txt.bak',
])

heading2('Real-World One-Liners You Will Use at Work')
code_block([
    '# Count ERROR lines in the last 200 lines of a running log',
    'tail -200 app.log | grep "ERROR" | wc -l',
    '',
    '# Find the 10 most frequent error messages',
    'grep "ERROR" app.log | sort | uniq -c | sort -rn | head -10',
    '',
    '# Count how many HTTP 500 errors in nginx access log',
    'grep " 500 " /var/log/nginx/access.log | wc -l',
    '',
    '# Find top 5 IP addresses hitting your server',
    'awk \'{print $1}\' /var/log/nginx/access.log | sort | uniq -c | sort -rn | head -5',
    '',
    '# Replace database host in application.properties (for deployment)',
    'sed -i "s/localhost:3306/prod-db.us-east-1.rds.amazonaws.com:3306/g" application.properties',
    '',
    '# Check disk usage of all directories in /var/log, sorted largest first',
    'du -sh /var/log/* | sort -rh | head -10',
])

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 10 — DSA BLOCK
# ══════════════════════════════════════════════════════════════════════════════
page_break()
heading1('Section 10 — DSA Block: Two Sum  (12:00 PM – 1:00 PM)')

body('Platform: LeetCode Problem #1 — Two Sum  |  Difficulty: Easy  |  Pattern: HashMap / Complement Lookup')
body('Rule: 30 minutes to solve on your own. If stuck at 30 min, read the approach. NEVER copy-paste code.')

youtube_box([
    ('Two Sum — Brute Force to HashMap Optimal', 'NeetCode',      'https://www.youtube.com/watch?v=KLlXCFG5TnA'),
    ('Big-O Notation — O(n) vs O(n²) Explained', 'YouTube Search', 'https://www.youtube.com/results?search_query=big+o+notation+o+n+vs+o+n+squared+explained'),
])

important_box(
    'Watch the NeetCode video ONLY AFTER you attempt the problem yourself. '
    'If you watch first, you are not practising — you are observing. '
    'The interview will not have a walkthrough playing beside you.',
    label='HOW TO USE THESE VIDEOS'
)

heading2('The Problem Statement')
code_block([
    'Given an array of integers nums and an integer target,',
    'return the INDICES of the two numbers that add up to target.',
    '',
    'Rules:',
    '  - Exactly one solution is guaranteed to exist.',
    '  - You cannot use the same element twice.',
    '  - Return the indices in any order.',
    '',
    'Example 1:  nums = [2, 7, 11, 15],  target = 9   →  Output: [0, 1]',
    '            nums[0] + nums[1] = 2 + 7 = 9  ✓',
    '',
    'Example 2:  nums = [3, 2, 4],  target = 6   →  Output: [1, 2]',
    '            nums[1] + nums[2] = 2 + 4 = 6  ✓',
    '',
    'Example 3:  nums = [3, 3],  target = 6   →  Output: [0, 1]',
    '            nums[0] + nums[1] = 3 + 3 = 6  ✓ (using different indices)',
])

heading2('Step 1 — Brute Force (Always Write This First in Interviews)')
code_block([
    'public int[] twoSum(int[] nums, int target) {',
    '    for (int i = 0; i < nums.length; i++) {',
    '        for (int j = i + 1; j < nums.length; j++) {  // j starts AFTER i to avoid reuse',
    '            if (nums[i] + nums[j] == target) {',
    '                return new int[]{i, j};',
    '            }',
    '        }',
    '    }',
    '    return new int[]{};',
    '}',
    '',
    '// Time complexity:  O(n²)  — nested loops, every pair is checked',
    '// Space complexity: O(1)   — no extra data structures',
])

heading2('Step 2 — Optimal Solution: HashMap Complement Lookup')
body(
    'The key insight: for every number x at index i, we need complement = target - x. '
    'If we store every number we have already seen (and its index) in a HashMap, we can check '
    'for the complement in O(1) — constant time — converting O(n²) brute force into O(n) one pass.'
)
code_block([
    'public int[] twoSum(int[] nums, int target) {',
    '    // HashMap: key = number value, value = its index in the array',
    '    Map<Integer, Integer> seen = new HashMap<>();',
    '',
    '    for (int i = 0; i < nums.length; i++) {',
    '        int complement = target - nums[i];   // What number do we NEED?',
    '',
    '        if (seen.containsKey(complement)) {  // Have we SEEN that number before?',
    '            return new int[]{seen.get(complement), i};',
    '        }',
    '',
    '        seen.put(nums[i], i);   // Store current number → index mapping',
    '    }',
    '',
    '    return new int[]{};',
    '}',
    '',
    '// Time complexity:  O(n)   — single pass, HashMap ops are O(1) average',
    '// Space complexity: O(n)   — HashMap stores at most n entries',
])

heading2('Full Trace Through Both Examples')
code_block([
    '# Example 1: nums = [2, 7, 11, 15],  target = 9',
    '',
    'i=0: current=2,  complement=9-2=7.   seen={}.     7 not in seen. Add 2→0. seen={2:0}',
    'i=1: current=7,  complement=9-7=2.   seen={2:0}.  2 IS in seen! Return [seen[2], 1] = [0, 1] ✓',
    '',
    '# Example 2: nums = [3, 2, 4],  target = 6',
    '',
    'i=0: current=3,  complement=6-3=3.   seen={}.     3 not in seen. Add 3→0. seen={3:0}',
    'i=1: current=2,  complement=6-2=4.   seen={3:0}.  4 not in seen. Add 2→1. seen={3:0, 2:1}',
    'i=2: current=4,  complement=6-4=2.   seen={3:0, 2:1}.  2 IS in seen! Return [1, 2] ✓',
    '',
    '# Example 3: nums = [3, 3],  target = 6',
    '',
    'i=0: current=3,  complement=6-3=3.   seen={}.     3 not in seen. Add 3→0. seen={3:0}',
    'i=1: current=3,  complement=6-3=3.   seen={3:0}.  3 IS in seen! Return [seen[3], 1] = [0, 1] ✓',
    '     Note: seen[3]=0, which is a DIFFERENT index from i=1. Uses different elements. ✓',
])

heading2('What to Say Out Loud During the Interview')
body(
    '"Let me start with the brute force — nested loops, check every pair, O(n²) time and O(1) space. '
    'That is correct but slow. For n=10,000 elements that is 50 million comparisons. '
    'We can do better. For each number x, I need its complement — target minus x. '
    'If I use a HashMap to store numbers I have already seen, I can check for the complement '
    'in O(1) instead of scanning the whole array again. That gives me O(n) time — one pass — '
    'and O(n) space for the HashMap. Let me trace through example 1 to verify before coding..."'
)

tip_box(
    'The interviewer wants to see: (1) brute force spoken first, (2) recognition of the inefficiency, '
    '(3) the insight about HashMap lookups, (4) a trace to verify correctness, (5) complexity analysis. '
    'Silence during coding is bad. Think out loud throughout.',
    label='INTERVIEW STRATEGY'
)

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 11 — INTERVIEW Q&A
# ══════════════════════════════════════════════════════════════════════════════
page_break()
heading1('Section 11 — Interview Q&A (1:00 PM – 2:00 PM)')

important_box(
    'INSTRUCTIONS: Read each question. CLOSE this document. Write your answer on paper from memory. '
    'Open the doc and compare. Then say your answer OUT LOUD — speak it to the wall if needed. '
    'If you cannot say it fluently in 60 seconds, you do not know it yet. Keep practicing.',
    label='HOW TO USE THIS SECTION'
)

qa_block(
    'What does chmod 755 mean? Break down every digit.',
    'chmod 755 sets the permission to rwxr-xr-x. Each digit represents one category. '
    'The first digit 7 is for the OWNER: read(4) + write(2) + execute(1) = 7 — full access. '
    'The second digit 5 is for the GROUP: read(4) + execute(1) = 5 — can read and run, cannot modify. '
    'The third digit 5 is for OTHERS: same as group. '
    'chmod 755 is used for executable scripts and public binaries.'
)

qa_block(
    'How do you find a running process and kill it gracefully?',
    'I run ps aux | grep java to find the process. The second column is the PID. '
    'I use kill <PID> to send SIGTERM — a graceful shutdown request. The application can handle '
    'this signal: Spring Boot will stop accepting new connections, finish in-progress requests, '
    'close DB connections, and run @PreDestroy methods. I wait 30 seconds. '
    'If still running, I escalate to kill -9 <PID> which is SIGKILL — immediate and uncatchable, '
    'but risks leaving resources inconsistent. For port conflicts, lsof -i :8080 shows me '
    'which process has that port.'
)

qa_block(
    'What is the difference between > and >> in Linux?',
    '> redirects command output to a file and OVERWRITES the file — all previous content is lost. '
    '>> APPENDS to the end of the file without touching existing content. '
    'Getting this wrong destroys data. In DevOps I always use >> when logging command output: '
    '"mvn package >> build.log 2>&1" appends each build\'s output to the same log file. '
    'The 2>&1 merges stderr into stdout so errors also go to the log, not just to the terminal.'
)

qa_block(
    'What is the difference between SIGTERM and SIGKILL?',
    'SIGTERM (15) is a polite termination request. The process receives it, handles it, finishes work, '
    'closes connections, and shuts down cleanly. Spring Boot, for example, runs @PreDestroy hooks. '
    'SIGKILL (9) is different at the kernel level — it cannot be caught, blocked, or ignored. '
    'The kernel kills the process immediately. There is no cleanup. DB connections leak, '
    'open files may be partially written and corrupt, in-flight HTTP requests drop. '
    'Always SIGTERM first, SIGKILL only if the process is frozen and unresponsive.'
)

qa_block(
    'What is the /proc filesystem?',
    '/proc is a virtual filesystem generated by the kernel — nothing in it is stored on disk. '
    'It is the kernel\'s window into itself. /proc/<PID>/ is a directory for each running process '
    'containing its memory maps, open file descriptors, environment variables, and command line. '
    '/proc/cpuinfo shows CPU details, /proc/meminfo shows memory stats — exactly what the free '
    'command reads. Tools like top, ps, and htop get all their data from /proc. '
    'This is the Linux "everything is a file" philosophy applied to kernel internals.'
)

qa_block(
    'How do you check which process is using port 8080?',
    'I use lsof -i :8080. lsof means "list open files" and in Linux, network sockets are files. '
    'The output shows the process name, PID, user, and connection state. '
    'Alternatively, ss -tlnp | grep 8080 is the modern method — ss replaced netstat. '
    'With the PID I can then kill the process or investigate further with ps aux | grep <PID>. '
    'This comes up frequently in Spring Boot development: "address already in use" error '
    'means a previous instance is still running on 8080.'
)

qa_block(
    'How do you check if a server is running out of disk space or memory?',
    'For disk: df -h shows disk usage of all filesystems in human-readable format. '
    'I look for any filesystem above 80% usage. To find what is consuming space: '
    'du -sh /var/log/* | sort -rh | head -10 finds the largest directories. '
    'For memory: free -h shows total, used, and available RAM. The "available" column is '
    'what your application can actually use — not just "free". '
    'On AWS I set CloudWatch alarms on disk and memory metrics to alert before we reach 80%.'
)

qa_block(
    'Explain the Two Sum HashMap approach and its time complexity.',
    'For each element x at index i, I calculate the complement = target - x. '
    'I check a HashMap to see if the complement was already seen in a previous iteration. '
    'If yes, I return [index_of_complement, i]. If no, I store x → i in the HashMap and continue. '
    'This converts the O(n²) brute force (nested loops checking every pair) into O(n) — a single pass '
    'through the array. HashMap containsKey and get operations are O(1) average time complexity. '
    'Space complexity is O(n) since the HashMap stores at most n entries.'
)

qa_block(
    'How does sudo work and why do we need it?',
    'sudo (superuser do) allows a permitted user to run commands as root without switching '
    'to the root account. The system checks /etc/sudoers to see if your user is allowed, '
    'then asks for YOUR own password. This is safer than logging in as root directly because: '
    '(1) there is an audit trail of every sudo command in /var/log/auth.log, '
    '(2) you only elevate for one command at a time, not an entire session, '
    'and (3) your own account permissions remain limited for all other work. '
    'On AWS EC2 with Amazon Linux, ec2-user has passwordless sudo configured by default.'
)

qa_block(
    'What is a zombie process? What causes it and should you be worried?',
    'A zombie process is a process that has finished executing but whose entry remains in the '
    'process table because its parent process has not yet called wait() to collect its exit status. '
    'In ps aux output, zombies appear with STAT code Z: "Z = defunct". '
    'Zombies do not consume CPU or memory — they only occupy a PID slot in the process table. '
    'A few zombies are normal and harmless. Many zombies indicate a bug in the parent process '
    'that is not reaping its children properly. You cannot kill a zombie with kill -9 because '
    'it is already dead — only killing the parent process (which triggers cleanup) removes them. '
    'In interviews, mention that seeing zombie accumulation signals a parent process bug, '
    'not a problem with the zombie process itself.'
)

qa_block(
    'What is the difference between a hard link and a symbolic link in Linux?',
    'A hard link (ln source dest) creates a second directory entry pointing to the SAME inode — '
    'the same physical disk blocks. Deleting the original filename does not affect the hard link '
    'because the data persists until ALL hard links to that inode are removed. '
    'Hard links cannot cross filesystem boundaries and cannot point to directories. '
    'A symbolic link (ln -s source dest) stores the PATH to the target as text. '
    'If the target is deleted or moved, the symlink becomes a dangling broken link. '
    'Symlinks can cross filesystems and can point to directories. '
    'In ls -la, symlinks show "l" at position 1 and display the target with ->: '
    'lrwxrwxrwx  java -> /opt/java-21/bin/java. '
    'I use symlinks for Java version switching and config file aliases on servers.'
)

qa_block(
    'What does 2>&1 mean in a Linux command?',
    'Every process has three standard file descriptors: stdin (0), stdout (1), stderr (2). '
    '2>&1 means redirect file descriptor 2 (stderr) to wherever file descriptor 1 (stdout) currently points. '
    'In: nohup java -jar app.jar > app.log 2>&1 & '
    '-- the > app.log redirects stdout to the file first, '
    'then 2>&1 redirects stderr to wherever stdout now points, which is app.log. '
    'Both normal output and error messages (like stack traces) are captured in one file. '
    'The order matters: 2>&1 must come AFTER the > redirection. '
    'Without 2>&1, errors still print to the terminal even though normal output goes to the file.'
)

qa_block(
    'What is the difference between an absolute path and a relative path?',
    'An absolute path starts from the root directory (/) and gives the complete location of a file '
    '-- for example /etc/nginx/nginx.conf. It works correctly from any directory on the system. '
    'A relative path is relative to your current working directory '
    '-- for example ./config.properties or ../logs/app.log. '
    'It only works correctly if you are in the right directory at the time. '
    '. means the current directory, .. means the parent directory, ~ means your home directory. '
    'In shell scripts I always use absolute paths because the script may be called from any location '
    '(cron, systemd, CI/CD pipelines) and a relative path would point to the wrong file. '
    'Interactive navigation on the command line typically uses relative paths for speed.'
)

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 12 — PRACTICAL BLOCK
# ══════════════════════════════════════════════════════════════════════════════
page_break()
heading1('Section 12 — Practical Block (2:10 PM – 4:10 PM)')

important_box(
    'TYPE every command — do not copy-paste. Your hands must build muscle memory. '
    'When a command fails, READ the error message completely before trying again. '
    'Error messages in Linux are usually accurate and helpful.',
    label='HANDS-ON RULE'
)

heading2('Lab 1: Install Git Bash and First Commands  (2:10 – 2:30 PM)')
bullet('Go to: https://git-scm.com/download/win')
bullet('Download the Windows installer. Run with all default options.')
bullet('After install: right-click on Desktop → "Git Bash Here"')
bullet('This gives you a real Linux shell on Windows — bash, grep, find, awk, sed, all included.')
body('Type every command below — do not skip:')
code_block([
    '# Verify your shell',
    'echo $SHELL                # Should show: /usr/bin/bash or similar',
    'uname -a                   # Show system info',
    'whoami                     # Show current username',
    'pwd                        # Where are you?',
    'ls -la                     # What is here?',
    'cd ~                       # Go to home',
    '',
    '# Create your day1 practice folder structure',
    'mkdir -p ~/day1_practice/logs',
    'mkdir -p ~/day1_practice/scripts',
    'mkdir -p ~/day1_practice/configs',
    'cd ~/day1_practice',
    'ls -la                     # Verify 3 subdirs were created',
    '',
    '# Create a realistic Java Spring Boot project layout',
    '# (This is the standard Maven project structure used at every Java company)',
    'mkdir -p ~/projects/myapp/src/main/java/com/example',
    'mkdir -p ~/projects/myapp/src/main/resources',
    'mkdir -p ~/projects/myapp/src/test/java/com/example',
    'mkdir -p ~/projects/myapp/target',
    '',
    '# Navigate and verify the structure',
    'cd ~/projects/myapp',
    'find . -type d              # List ALL directories -- you should see the Java tree',
    'ls -R                      # Recursive listing of all nested directories',
    'pwd                        # Confirm where you are (absolute path)',
])

heading2('Lab 2: File Operations and Content Creation  (2:30 – 3:00 PM)')
code_block([
    'cd ~/day1_practice/configs',
    '',
    '# Create application.properties step by step (using >> to append)',
    'echo "server.port=8080" > application.properties',
    'echo "server.servlet.context-path=/api" >> application.properties',
    'echo "" >> application.properties',
    'echo "# Database config" >> application.properties',
    'echo "spring.datasource.url=jdbc:mysql://localhost:3306/traveldb" >> application.properties',
    'echo "spring.datasource.username=appuser" >> application.properties',
    'echo "spring.datasource.password=changeme123" >> application.properties',
    'echo "spring.datasource.hikari.maximum-pool-size=10" >> application.properties',
    'echo "" >> application.properties',
    'echo "# Redis config" >> application.properties',
    'echo "spring.redis.host=localhost" >> application.properties',
    'echo "spring.redis.port=6379" >> application.properties',
    '',
    '# Read it multiple ways',
    'cat application.properties            # Full file',
    'cat -n application.properties         # With line numbers',
    'head -5 application.properties        # First 5 lines',
    'tail -4 application.properties        # Last 4 lines',
    'wc -l application.properties          # Count lines',
    '',
    '# Search inside it',
    'grep "datasource" application.properties',
    'grep "port" application.properties',
    'grep -n "spring" application.properties     # With line numbers',
    'grep -v "#" application.properties          # Lines that are NOT comments',
    '',
    '# Create placeholder Java files to practice find',
    'touch ~/projects/myapp/src/main/java/com/example/Application.java',
    'touch ~/projects/myapp/src/main/java/com/example/BookingService.java',
    'touch ~/projects/myapp/src/main/java/com/example/PaymentController.java',
    'touch ~/projects/myapp/src/test/java/com/example/BookingServiceTest.java',
    'touch ~/projects/myapp/src/main/resources/application.properties',
    '',
    '# Practice find -- this is directly from your 30-day study plan',
    'find ~/projects -name "*.java"              # Find ALL .java files recursively',
    'find ~/projects -name "*Service*"           # Files containing "Service" in name',
    'find ~/projects -name "*Test*"              # Find test files specifically',
    'find ~/projects -type d                     # Find all DIRECTORIES (no files)',
    'find ~/projects -name "*.java" | wc -l     # Count how many .java files exist',
    'find ~/projects -type f -name "*.java"     # Only regular files (not dirs) matching *.java',
])

heading2('Lab 3: Simulate an Application Log File  (3:00 – 3:20 PM)')
code_block([
    'cd ~/day1_practice/logs',
    '',
    "# Create a realistic log file (note: single quotes inside the heredoc are fine)",
    "cat > app.log << 'ENDOFLOG'",
    '[INFO]  2026-05-17 09:00:01.234 Application started on port 8080',
    '[INFO]  2026-05-17 09:00:02.001 Connected to MySQL: localhost:3306/traveldb',
    '[INFO]  2026-05-17 09:00:02.500 Redis connection established: localhost:6379',
    '[INFO]  2026-05-17 09:01:15.100 POST /api/auth/login - User U001 authenticated',
    '[ERROR] 2026-05-17 09:02:01.300 NullPointerException in BookingService.createBooking():142',
    '[WARN]  2026-05-17 09:03:00.500 Redis response time 285ms (threshold: 200ms)',
    '[ERROR] 2026-05-17 09:03:45.600 Database connection timeout after 30000ms',
    '[INFO]  2026-05-17 09:04:00.700 User U003 booked flight AI-302 successfully',
    '[ERROR] 2026-05-17 09:05:12.800 java.lang.OutOfMemoryError: GC overhead limit exceeded',
    '[WARN]  2026-05-17 09:05:30.900 Circuit breaker OPEN for PaymentService',
    '[ERROR] 2026-05-17 09:07:14.100 Connection refused: PaymentGateway at 10.0.0.45:443',
    '[INFO]  2026-05-17 09:08:00.200 Scheduled job: cleanup expired sessions - 47 removed',
    'ENDOFLOG',
    '',
    '# Verify creation',
    'wc -l app.log',
    'cat app.log',
    '',
    '# Now practice searching',
    'grep "ERROR" app.log                           # All error lines',
    'grep -c "ERROR" app.log                        # Count of errors',
    'grep -n "ERROR" app.log                        # Errors with line numbers',
    'grep -i "warn" app.log                         # Case insensitive',
    'grep -v "INFO" app.log                         # Everything except INFO',
    'grep -E "ERROR|WARN" app.log                   # ERROR or WARN',
])

heading2('Lab 4: Permissions Practice  (3:20 – 3:40 PM)')
code_block([
    'cd ~/day1_practice/scripts',
    '',
    '# Create a deployment script',
    "cat > deploy.sh << 'EOF'",
    '#!/bin/bash',
    'echo "============================================"',
    'echo " DEPLOYMENT SCRIPT"',
    'echo " User:      $(whoami)"',
    'echo " Directory: $(pwd)"',
    'echo " Date:      $(date)"',
    'echo "============================================"',
    'echo "Step 1: Checking disk space..."',
    'df -h .',
    'echo "Step 2: Script complete!"',
    'EOF',
    '',
    '# Check permissions — should be rw-r--r-- (644) by default',
    'ls -la deploy.sh',
    '',
    '# Try to run it — will fail: Permission denied',
    './deploy.sh',
    '',
    '# Add execute permission and try again',
    'chmod +x deploy.sh',
    'ls -la deploy.sh                    # Now shows rwxr-xr-x',
    './deploy.sh',
    '',
    '# Practice different permission levels',
    'touch private.key    ; chmod 600 private.key   ; ls -la private.key',
    'touch public.txt     ; chmod 644 public.txt    ; ls -la public.txt',
    'touch team_script.sh ; chmod 750 team_script.sh ; ls -la team_script.sh',
    'ls -la                              # See all at once',
])

heading2('Lab 5: Piping Challenges  (3:40 – 4:10 PM)')
body('Attempt each challenge yourself BEFORE looking at the solution. Write your command on paper first.')

heading3('Challenge 1 — Count how many WARN and ERROR lines exist in total')
body('Your command: _______________________________________')
code_block([
    '# Solution:',
    'grep -E "WARN|ERROR" ~/day1_practice/logs/app.log | wc -l',
])

heading3('Challenge 2 — Show only the time portion (HH:MM:SS) of ERROR lines')
body('Your command: _______________________________________')
code_block([
    '# Solution:',
    'grep "ERROR" ~/day1_practice/logs/app.log | awk \'{print $2}\'',
])

heading3('Challenge 3 — Count occurrences of each log level [INFO], [WARN], [ERROR]')
body('Your command: _______________________________________')
code_block([
    '# Solution:',
    'grep -o "\\[INFO\\]\\|\\[WARN\\]\\|\\[ERROR\\]" ~/day1_practice/logs/app.log | sort | uniq -c | sort -rn',
])

heading3('Challenge 4 — Find all lines that mention a specific user (U003)')
body('Your command: _______________________________________')
code_block([
    '# Solution:',
    'grep "U003" ~/day1_practice/logs/app.log',
])

heading3('Challenge 5 — Replace "localhost" with "prod-db.internal" in application.properties')
body('Your command: _______________________________________')
code_block([
    '# Solution (preview only — does NOT edit the file):',
    'sed "s/localhost/prod-db.internal/g" ~/day1_practice/configs/application.properties',
    '',
    '# Solution (edit the file IN PLACE and create a backup):',
    'sed -i.bak "s/localhost/prod-db.internal/g" ~/day1_practice/configs/application.properties',
    'cat ~/day1_practice/configs/application.properties  # Verify the change',
    'cat ~/day1_practice/configs/application.properties.bak  # Original is here',
])

heading3('Challenge 6 — Live Log Tailing (open two Git Bash windows)')
code_block([
    '# TERMINAL 1 — Start watching the log live:',
    'tail -f ~/day1_practice/logs/app.log',
    '',
    '# TERMINAL 2 — Add new entries (watch them appear instantly in Terminal 1):',
    'echo "[INFO]  2026-05-17 10:00:00 New booking created for user U005" >> ~/day1_practice/logs/app.log',
    'echo "[ERROR] 2026-05-17 10:00:01 Payment failed: card declined for B9999" >> ~/day1_practice/logs/app.log',
    'echo "[WARN]  2026-05-17 10:00:03 High memory usage: 85% of heap used" >> ~/day1_practice/logs/app.log',
    '',
    '# In Terminal 1: press Ctrl+C to stop tailing when done',
])

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 13 — END OF DAY
# ══════════════════════════════════════════════════════════════════════════════
page_break()
heading1('Section 13 — End of Day: Notes + Checklist (4:10 – 4:30 PM)')

heading2('Write in Your Own Words — NOT Copied From This Document')
body('Open a notebook or Notion page. Answer these 13 questions without looking at anything:')
bullet('1. What is Linux? Why does every cloud server run Linux? (3 sentences)')
bullet('2. List the 5 most important directories and what is in each one.')
bullet('3. What is the difference between > and >>? Give a real example.')
bullet('4. What does chmod 755 mean? Break down every digit.')
bullet('5. How do you find a process and kill it gracefully? What command and why?')
bullet('6. What is the difference between SIGTERM and SIGKILL?')
bullet('7. Draw -rwxr-xr-- and label every character and its numeric value.')
bullet('8. How do you check disk space? Memory? (commands)')
bullet('9. What is the HashMap approach for Two Sum? What is its time complexity and why?')
bullet('10. What is /proc? Why is it special?')
bullet('11. What is the difference between an absolute path and a relative path?')
bullet('12. What is the difference between a hard link and a symbolic link? (ln vs ln -s)')
bullet('13. What does 2>&1 mean and why does order matter in "cmd > out.txt 2>&1"?')
body('If you cannot answer any of these, re-read that section before sleeping.')

heading2('Day 1 Completion Checklist')
simple_table(
    ['Checkpoint', 'Done?'],
    [
        ['I navigated the Linux file system without looking at notes',                  '☐'],
        ['I know the difference between absolute path and relative path',               '☐'],
        ['I understand hard link vs symbolic link and can create a symlink (ln -s)',    '☐'],
        ['I know what stdin, stdout, stderr are and what 2>&1 means',                  '☐'],
        ['I created files, wrote to them, appended, and read them back',               '☐'],
        ['I understand chmod 755 and can explain it in 60 seconds out loud',           '☐'],
        ['I found a process by name and understand how to kill it',                    '☐'],
        ['I know SIGTERM vs SIGKILL and when to use each',                             '☐'],
        ['I used df -h and free -h and understand what the output means',              '☐'],
        ['I created the Java Maven project structure with mkdir -p',                   '☐'],
        ['I used find -name "*.java" to search for files',                             '☐'],
        ['I completed all 6 practical labs',                                           '☐'],
        ['I solved Two Sum with the HashMap approach on my own',                       '☐'],
        ['I traced through both examples manually on paper',                           '☐'],
        ['I answered all 13 interview Q&A out loud from memory (not reading)',         '☐'],
        ['I watched the NeetCode Two Sum video after solving it myself',               '☐'],
        ['I wrote my notes in my own words',                                           '☐'],
        ['I read the Day 2 plan and know what tomorrow covers',                        '☐'],
    ],
    header_bg='1A233E'
)

heading2('Day 2 Preview — What You Will Learn Tomorrow')
simple_table(
    ['Topic', 'Key Things You Will Learn'],
    [
        ['SSH Keys',              'How public/private key auth works, ssh-keygen, authorized_keys, known_hosts'],
        ['curl and wget',         'Make HTTP requests from terminal, test APIs, download files'],
        ['Environment Variables', '$PATH, export, .bashrc vs .bash_profile, why Docker needs them'],
        ['systemctl',             'start, stop, restart, enable, status of services — production standard'],
        ['Package Managers',      'apt (Ubuntu) and yum/dnf (Amazon Linux) — install, update, remove'],
        ['DSA: Day 2 Problem',    'Best Time to Buy and Sell Stock (LeetCode #121) — sliding window'],
    ]
)

tip_box(
    'Nothing to install before Day 2. Git Bash from today is sufficient. '
    'Day ends NOW at 4:30 PM. Rest, eat, sleep. Your brain will consolidate all of today '
    'during sleep — this is science, not a suggestion.',
    label='DAY IS DONE AT 4:30 PM'
)

divider()

heading2('Quick Reference Card')
code_block([
    '# NAVIGATION           # FILES                  # PERMISSIONS',
    'pwd                    touch file.txt           chmod 755 script.sh',
    'ls -la                 cat file.txt             chmod 600 key.pem',
    'cd ~                   less file.txt            chmod +x file',
    'cd ..                  tail -f log.txt          chown user file',
    '',
    '# SEARCH               # PROCESSES              # SYSTEM INFO',
    'grep -n "text" file    ps aux                   df -h',
    'grep -r "text" dir/    ps aux | grep java       free -h',
    'find . -name "*.log"   kill PID                 uname -a',
    'grep -E "A|B" file     kill -9 PID              uptime',
    '',
    '# REDIRECTION          # PORT CHECK             # TEXT OPS',
    'cmd > file (overwrite) lsof -i :8080            wc -l file',
    'cmd >> file (append)   ss -tlnp | grep 8080     sort | uniq -c',
    'cmd 2>&1               nohup cmd &              awk \'{print $2}\'',
    'cmd > /dev/null        history | grep cmd       sed -i "s/a/b/g" f',
])

divider()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(6)
r = p.add_run('Day 1 of 30  ·  30-Day Interview Preparation  ·  Target: 24 LPA  ·  Kanagarajan M')
r.font.size = Pt(9); r.font.color.rgb = GREY_TEXT; r.font.name = 'Calibri'

# ── Save ──────────────────────────────────────────────────────────────────────
path = r'C:\Users\Kanag\kanagarajan_interview_preparation_may15_2026\Day_01_Linux_Basics_FINAL.docx'
doc.save(path)
print(f'SUCCESS: {path}')

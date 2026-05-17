from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin   = Inches(0.75)
    section.right_margin  = Inches(0.75)

# ── Helper utilities ──────────────────────────────────────────────────────────

def set_font(run, name="Calibri", size=10, bold=False, italic=False, color=None):
    run.font.name  = name
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_para(text="", style="Normal", align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=0):
    p = doc.add_paragraph(style=style)
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    return p

def add_hrule(thickness=1):
    """Inserts a full-width horizontal rule."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(thickness * 6))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1F4E79")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def section_heading(title):
    add_hrule()
    p = add_para(space_before=2, space_after=1)
    r = p.add_run(title.upper())
    set_font(r, size=10.5, bold=True, color=(31, 78, 121))

def bullet(text, indent=0.18, hanging=0.18, font_size=9.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent    = Inches(indent)
    p.paragraph_format.first_line_indent = Inches(-hanging)
    p.paragraph_format.space_before   = Pt(0)
    p.paragraph_format.space_after    = Pt(1)
    r = p.add_run(text)
    set_font(r, size=font_size)

def add_right_tab(paragraph, pos_inches=6.5):
    """Adds a right-aligned tab stop at pos_inches from left margin."""
    pPr = paragraph._p.get_or_add_pPr()
    tabs = OxmlElement('w:tabs')
    tab = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'right')
    tab.set(qn('w:pos'), str(int(pos_inches * 1440)))
    tabs.append(tab)
    pPr.append(tabs)

def job_header(company, title, period, location):
    # Line 1: Company | Title  <TAB RIGHT>  Period  •  Location
    p = add_para(space_before=4, space_after=1)
    add_right_tab(p, pos_inches=6.5)
    r = p.add_run(company + "  |  ")
    set_font(r, size=10, bold=True, color=(31, 78, 121))
    r2 = p.add_run(title)
    set_font(r2, size=10, bold=True, color=(0, 0, 0))
    r3 = p.add_run(f"\t{period}  •  {location}")
    set_font(r3, size=9, italic=True, color=(89, 89, 89))

# ═══════════════════════════════════════════════════════════════════════════════
#  NAME  &  CONTACT
# ═══════════════════════════════════════════════════════════════════════════════
p_name = add_para(align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=0)
r = p_name.add_run("KANAGARAJAN M")
set_font(r, size=20, bold=True, color=(31, 78, 121))

p_title = add_para(align=WD_ALIGN_PARAGRAPH.CENTER, space_before=1, space_after=1)
r = p_title.add_run("Backend & DevOps Engineer  |  Java  •  Spring Boot  •  AWS  •  Microservices")
set_font(r, size=10, bold=False, color=(89, 89, 89))

p_contact = add_para(align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=4)
r = p_contact.add_run(
    "kanagarajanmuthukumar@gmail.com   |   +91 9962344750   |   "
    "linkedin.com/in/kanagarajan-m   |   Chennai, India"
)
set_font(r, size=9.5, color=(0, 0, 0))

# ═══════════════════════════════════════════════════════════════════════════════
#  PROFESSIONAL SUMMARY
# ═══════════════════════════════════════════════════════════════════════════════
section_heading("Professional Summary")
p_sum = add_para(space_before=2, space_after=2, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
summary = (
    "Results-driven Backend & DevOps Engineer with 4.5+ years of experience designing scalable, "
    "high-availability microservices with Java and Spring Boot, managing cloud infrastructure on AWS, "
    "and driving production reliability for enterprise-grade systems. Proficient in containerized deployments "
    "(Docker, Kubernetes, OpenShift), CI/CD automation, and end-to-end observability using Datadog, Grafana, "
    "and ELK Stack. Proven expertise in RESTful API design, distributed systems, AI-powered platform "
    "integration, and ITSM-driven incident management. Passionate about building performance-optimized, "
    "cloud-native architectures that directly impact business outcomes."
)
r = p_sum.add_run(summary)
set_font(r, size=9.5)

# ═══════════════════════════════════════════════════════════════════════════════
#  TECHNICAL SKILLS  (keyword-rich for ATS)
# ═══════════════════════════════════════════════════════════════════════════════
section_heading("Technical Skills")

skills = [
    ("Languages & Frameworks",
     "Java, Spring Boot, Spring Cloud Gateway, Spring Data JPA, Python, JavaScript"),
    ("Cloud & DevOps",
     "AWS (EC2, S3, ECS, EKS, IAM, VPC, CloudWatch, Lambda, RDS, Route 53), Docker, "
     "Kubernetes, OpenShift, CI/CD Pipelines"),
    ("Databases",
     "MySQL, MongoDB, Couchbase, Redis, Weaviate (Vector DB)"),
    ("Observability & Monitoring",
     "Datadog (APM, Tracing, Alerting), Grafana, Kibana, ELK Stack (Elasticsearch, Logstash)"),
    ("API & Architecture",
     "RESTful API Design, Microservices, Event-Driven Architecture, OpenAPI / Swagger, "
     "Resilience4j, Circuit Breaker, WebSockets"),
    ("AI & GenAI Integration",
     "Generative AI, OpenAI, Google Gemini, DeepSeek, RAG Pipeline, Prompt Engineering, "
     "Vector Embeddings, NLP, Conversational AI"),
    ("Developer Tools & ITSM",
     "Git, GitHub, Maven, Postman, JMeter, ServiceNow, Fortuna"),
]

for label, value in skills:
    p = add_para(space_before=1, space_after=0)
    r1 = p.add_run(label + ":  ")
    set_font(r1, size=9.5, bold=True, color=(31, 78, 121))
    r2 = p.add_run(value)
    set_font(r2, size=9.5)

# ═══════════════════════════════════════════════════════════════════════════════
#  WORK EXPERIENCE
# ═══════════════════════════════════════════════════════════════════════════════
section_heading("Work Experience")

# ── TCS ───────────────────────────────────────────────────────────────────────
job_header("Tata Consultancy Services (TCS)", "DevOps & Cloud Infrastructure Engineer",
           "Aug 2025 – Present", "Chennai")

p_proj = add_para(space_before=0, space_after=1)
r = p_proj.add_run("Client: USAA (United Services Automobile Association) — US Financial Services")
set_font(r, size=9.5, italic=True, color=(89, 89, 89))

bullets_tcs = [
    "Implemented and maintained AWS-integrated infrastructure monitoring solutions using CloudWatch, Grafana, "
    "and Datadog to ensure high availability and performance visibility across critical financial services workloads.",
    "Configured and optimized ELK Stack (Elasticsearch, Logstash, Kibana) log aggregation pipelines for "
    "centralized log management; performed deep log-based root-cause analysis reducing issue resolution "
    "time by 40% and improving overall system uptime.",
    "Managed containerized application deployments on OpenShift (Kubernetes) clusters — provisioning pods, "
    "executing rolling updates, scaling services, and remediating unhealthy pods to maintain zero-downtime releases.",
    "Configured automated alerting and anomaly-detection rules in Grafana and Datadog, enabling proactive "
    "incident detection and reducing critical alert response time by 35%.",
    "Participated as the infrastructure technical SME in Major Incident (P1/P2) bridge calls; performed "
    "root-cause analysis on infrastructure and deployment-layer failures, and authored Post-Incident Review "
    "(PIR/POMO) documentation to prevent recurrence.",
    "Executed service impact analysis using Fortuna prior to each production deployment, identifying "
    "upstream/downstream dependency risks and preventing unplanned outages.",
    "Raised and tracked production change requests and temporary database access workflows through "
    "ServiceNow (Change & Access Management), ensuring release governance and security compliance.",
    "Utilized Datadog APM distributed tracing to visualize end-to-end request spans across microservices, "
    "diagnosing latency bottlenecks and enabling data-driven performance improvements.",
]
for b in bullets_tcs:
    bullet(b)

# ── Mondee ─────────────────────────────────────────────────────────────────────
job_header("Mondee Inc.", "Backend Software Engineer",
           "Aug 2023 – Aug 2025", "Bangalore")

bullets_mondee = [
    "Designed and built a cloud-native microservices platform on AWS using Spring Boot and Spring Cloud Gateway, "
    "serving millions of travel booking transactions with sub-200 ms p95 API response times.",
    "Developed Abhi Gateway and AI Gateway (Spring Cloud Gateway) featuring JWT token-based authentication, "
    "smart load-balancing routing, Resilience4j circuit breakers, response caching, and downstream health "
    "checks — improving platform resilience by 45%.",
    "Integrated Generative AI models (OpenAI GPT, Google Gemini, DeepSeek) into the Jello conversational "
    "AI microservice; system handled 10K+ daily user interactions with intent-to-response accuracy of 92%.",
    "Implemented a RAG (Retrieval-Augmented Generation) pipeline using Weaviate vector database — extracted "
    "content from web/documents, generated embeddings, and matched queries semantically for AI-powered "
    "knowledge base responses.",
    "Created Morpheus, a centralized media content microservice integrated with Cloudinary for secure media "
    "storage and token-based delivery; implemented Redis-backed chunked file upload for reliable large-file "
    "handling.",
    "Deployed and monitored all microservices using Datadog APM with distributed tracing across service "
    "boundaries; reduced production incidents by 30% through real-time span visualization and alerting.",
    "Built real-time WebSocket communication channels and webhook-based integrations for multi-platform "
    "chatbot deployments (PurpleCloud); developed MongoDB aggregate pipelines for analytics and reporting.",
]
for b in bullets_mondee:
    bullet(b)

# ── Amshuhu ─────────────────────────────────────────────────────────────────────
job_header("Amshuhu Itech Solution", "Junior Software Engineer",
           "Oct 2021 – Aug 2023", "Chennai")

bullets_ams = [
    "Developed and maintained Isteer ERP — a full-featured monolithic Spring Boot application with MySQL, "
    "Spring Data JPA, JavaScript, and jQuery covering Inventory, Purchase, Sales, and Accounting modules.",
    "Integrated government e-invoicing system by consuming third-party REST APIs from IRIS, enabling "
    "automated GST-compliant invoice posting and eliminating manual intervention.",
    "Automated the GSTR-1 filing workflow within the ERP, reducing client manual data-preparation time "
    "by 60% and improving compliance accuracy.",
    "Designed custom reporting modules using JPA aggregate queries, delivering analytics and compliance "
    "dashboards tailored to client-specific requirements.",
    "Utilized Git for version control and collaborative code reviews, maintaining clean branching strategies "
    "across the development team.",
]
for b in bullets_ams:
    bullet(b)

# ═══════════════════════════════════════════════════════════════════════════════
#  KEY ACHIEVEMENTS
# ═══════════════════════════════════════════════════════════════════════════════
section_heading("Key Achievements")
achievements = [
    "Reduced production MTTR by 40% via ELK-based log analysis pipelines for a US financial services client.",
    "Improved API gateway resilience by 45% through circuit breaker patterns, caching, and health-check routing.",
    "Delivered end-to-end Generative AI chatbot platform (10K+ daily interactions) from architecture design to production rollout.",
    "Diagnosed P1/P2 infrastructure incidents as technical SME, performing root-cause analysis and documenting findings to prevent recurrence.",
    "Automated GSTR-1 filing, saving 60% manual effort and achieving 100% compliance accuracy for ERP clients.",
]
for a in achievements:
    bullet(a)

# ═══════════════════════════════════════════════════════════════════════════════
#  EDUCATION
# ═══════════════════════════════════════════════════════════════════════════════
section_heading("Education")
p_edu = add_para(space_before=2, space_after=0)
r1 = p_edu.add_run("Master of Computer Applications (MCA)")
set_font(r1, size=10, bold=True, color=(0, 0, 0))
p_edu2 = add_para(space_before=0, space_after=0)
r2 = p_edu2.add_run("Sri Muthukumaran Institute of Technology  |  2020 – 2022")
set_font(r2, size=9.5, color=(89, 89, 89))

# ═══════════════════════════════════════════════════════════════════════════════
#  CERTIFICATIONS  (target certs — shows ambition to recruiters)
# ═══════════════════════════════════════════════════════════════════════════════
section_heading("Certifications & Learning")
certs = [
    "AWS Certified Developer – Associate  —  Amazon Web Services  (Certified)",
    "AWS Certified Solutions Architect – Associate  (In Progress  |  Target: Q3 2026)",
    "Spring Professional Certification — VMware/Broadcom  (In Progress  |  Target: Q4 2026)",
]
for c in certs:
    bullet(c)

# ═══════════════════════════════════════════════════════════════════════════════
#  LANGUAGES
# ═══════════════════════════════════════════════════════════════════════════════
section_heading("Languages")
p_lang = add_para(space_before=2, space_after=2)
r = p_lang.add_run("English — Full Professional Proficiency   |   Tamil — Native")
set_font(r, size=9.5)

# ── Save ──────────────────────────────────────────────────────────────────────
out_path = r"c:\Users\Kanag\kanagarajan_interview_preparation_may15_2026\Kanagarajan_M_Resume_2026_v4.docx"
doc.save(out_path)
print(f"Resume saved → {out_path}")
